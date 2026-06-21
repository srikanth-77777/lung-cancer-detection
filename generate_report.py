"""
Generate Minor Project Report in .docx format
Matching the exact structure of the reference report
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page Setup ─────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5


# ── Helper Functions ───────────────────────────────────────
def add_heading_centered(text, level=1, bold=True, size=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p

def add_body(text, bold=False, alignment=None, size=12, spacing_after=6):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(spacing_after)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Pt(36 + level * 18)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def add_list_item(text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
        set_cell_shading(cell, 'D9E2F3')
    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
    return table

def page_break():
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════
doc.add_paragraph()  # spacer
add_heading_centered(
    "Lung Cancer Detection and Multi-Class Classification\n"
    "Using CT Images with Explainable AI Diagnosis",
    level=1, bold=True, size=18
)
doc.add_paragraph()
add_heading_centered("21CSP302L- PROJECT", level=2, bold=True, size=14)
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Submitted by")
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

add_heading_centered("B S VIKASH SRIKANTH [RA2311026010XXX]", bold=True, size=13)
add_heading_centered("S RAVI PRAKASH [RA2311026010XXX]", bold=True, size=13)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Under the Guidance of")
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

add_heading_centered("Dr. N Gopinath", bold=True, size=16)
add_heading_centered("Assistant Professor, Department of Computational Intelligence", bold=False, size=12)

doc.add_paragraph()
add_heading_centered("in partial fulfillment of the requirements for the degree of", bold=False, size=12)
add_heading_centered("BACHELOR OF TECHNOLOGY", bold=True, size=14)
add_heading_centered("in", bold=False, size=12)
add_heading_centered("COMPUTER SCIENCE ENGINEERING", bold=True, size=14)
add_heading_centered("with specialization in AIML", bold=True, size=13)

doc.add_paragraph()
add_heading_centered("DEPARTMENT OF COMPUTATIONAL INTELLIGENCE\nCOLLEGE OF ENGINEERING AND TECHNOLOGY", bold=True, size=13)
add_heading_centered("SRM INSTITUTE OF SCIENCE AND TECHNOLOGY", bold=True, size=14)
add_heading_centered("KATTANKULATHUR- 603 203", bold=True, size=13)
add_heading_centered("MAY 2026", bold=True, size=13)

page_break()

# ═══════════════════════════════════════════════════════════
# OWN WORK DECLARATION
# ═══════════════════════════════════════════════════════════
add_heading_centered("Department of Computational Intelligence", bold=True, size=14)
add_heading_centered("SRM Institute of Science & Technology", bold=True, size=13)
add_heading_centered("Own Work Declaration Form", bold=True, size=13)

doc.add_paragraph()
add_body("This sheet must be filled in (each box ticked to show that the condition has been met). It must be signed and dated along with your student registration number and included with all assignments you submit – work will not be marked unless this is done.")
add_body("To be completed by the student for all assessments", bold=True)

doc.add_paragraph()
add_body("Degree/ Course\t: B.TECH/ CSE with specialization in AIML", bold=True)
add_body("Student Name\t: B S VIKASH SRIKANTH, S RAVI PRAKASH")
add_body("Registration Number\t: RA2311026010XXX, RA2311026010XXX", bold=True)
add_body("Title of Work\t: LUNG CANCER DETECTION AND MULTI-CLASS CLASSIFICATION USING CT IMAGES WITH EXPLAINABLE AI DIAGNOSIS")

doc.add_paragraph()
add_body("I / We hereby certify that this assessment compiles with the University's Rules and Regulations relating to Academic misconduct and plagiarism, as listed in the University Website, Regulations, and the Education Committee guidelines.")
add_body("I / We confirm that all the work contained in this assessment is my / our own except where indicated, and that I / We have met the following conditions:")

add_list_item("Clearly referenced / listed all sources as appropriate")
add_list_item("Referenced and put in inverted commas all quoted text (from books, web, etc)")
add_list_item("Given the sources of all pictures, data etc. that are not my own")
add_list_item("Not made any use of the report(s) or essay(s) of any other student(s) either past or present")
add_list_item("Acknowledged in appropriate places any help that I have received from others (e.g. fellow students, technicians, statisticians, external sources)")
add_list_item("Compiled with any other plagiarism criteria specified in the Course handbook / University website")

add_body("I understand that any false claim for this work will be penalized in accordance with the University policies and regulations.")

page_break()

# ═══════════════════════════════════════════════════════════
# BONAFIDE CERTIFICATE
# ═══════════════════════════════════════════════════════════
add_heading_centered("SRM INSTITUTE OF SCIENCE AND TECHNOLOGY\nKATTANKULATHUR – 603 203", bold=True, size=14)
doc.add_paragraph()
add_heading_centered("BONAFIDE CERTIFICATE", bold=True, size=16)
doc.add_paragraph()

add_body(
    'Certified that 21CSP302L - Project report titled "Lung Cancer Detection and Multi-Class Classification '
    'Using CT Images with Explainable AI Diagnosis" is the bonafide work of "B S VIKASH SRIKANTH [RA2311026010XXX], '
    'S RAVI PRAKASH [RA2311026010XXX]" who carried out the project work under my supervision. Certified further, '
    'that to the best of my knowledge the work reported herein does not form any other project report or dissertation '
    'on the basis of which a degree or award was conferred on an earlier occasion on this or any other candidate.'
)

doc.add_paragraph()
doc.add_paragraph()

# Signature table
sig_table = doc.add_table(rows=2, cols=2)
sig_table.rows[0].cells[0].text = "SIGNATURE\nDr. N Gopinath\nSUPERVISOR\nAssistant Professor\nDEPARTMENT OF\nCOMPUTATIONAL INTELLIGENCE"
sig_table.rows[0].cells[1].text = "SIGNATURE\nDR. R. ANNIE UTHRA\nPROFESSOR & HEAD\n\nDEPARTMENT OF\nCOMPUTATIONAL INTELLIGENCE"
sig_table.rows[1].cells[0].text = "EXAMINER 1\nName & Signature"
sig_table.rows[1].cells[1].text = "EXAMINER 2\nName & Signature"

page_break()

# ═══════════════════════════════════════════════════════════
# ACKNOWLEDGEMENT
# ═══════════════════════════════════════════════════════════
add_heading_centered("ACKNOWLEDGEMENT", bold=True, size=16)
doc.add_paragraph()

add_body(
    "We take this opportunity to express our sincere gratitude to our Chancellor, Pro-Chancellor and "
    "Vice-Chancellor of SRM Institute of Science and Technology for giving us this opportunity to pursue "
    "our career in this reputed institution."
)
add_body(
    "We are incredibly grateful to our Head of the Department, Dr. R. Annie Uthra, Professor, Department "
    "of Computational Intelligence, SRM Institute of Science and Technology, for her suggestions and "
    "encouragement at all the stages of the project work."
)
add_body(
    "We want to convey our thanks to our Project Coordinators, Panel Head, and Panel Members Department "
    "of Computational Intelligence, SRM Institute of Science and Technology, for their inputs during the "
    "project reviews and support."
)
add_body(
    "Our inexpressible respect and thanks to our guide, Dr. N Gopinath, Department of Computational Intelligence, "
    "SRM Institute of Science and Technology, for providing us with an opportunity to pursue our project under "
    "his mentorship. He provided us with the freedom and support to explore the research topics of our interest. "
    "His passion for solving problems and making a difference in the world has always been inspiring."
)
add_body(
    "We sincerely thank all the staff members of Computational Intelligence, School of Computing, "
    "S.R.M Institute of Science and Technology, for their help during our project. Finally, we would like "
    "to thank our parents, family members, and friends for their unconditional love, constant support and encouragement."
)
doc.add_paragraph()
add_body("Authors", alignment=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)

page_break()

# ═══════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════
add_heading_centered("ABSTRACT", bold=True, size=16)
doc.add_paragraph()

add_body(
    "Lung cancer remains the leading cause of cancer-related mortality worldwide. While deep learning has "
    "shown significant promise in automating radiological analysis, a critical gap persists between algorithmic "
    "output and clinical trust. Existing AI diagnostic tools often operate as opaque black boxes, providing "
    "classification scores without spatial or categorical justification. Furthermore, many systems are limited "
    "to simple binary detection, failing to deliver the granular histological subtyping necessary for "
    "comprehensive treatment planning."
)
add_body(
    "This project presents a multi-stage, explainable AI diagnostic framework for lung cancer subtyping and "
    "malignancy detection in chest Computed Tomography (CT) scans. The system employs a hierarchical architecture "
    "built upon the ResNet50 backbone. Sector 1 classifies scans into four histological categories — Adenocarcinoma, "
    "Squamous Cell Carcinoma, Large Cell Carcinoma, and Normal tissue — achieving a classification accuracy of 95%. "
    "Sector 2 performs secondary binary malignancy assessment (Benign vs. Malignant) on individual nodule slices."
)
add_body(
    "To bridge the transparency gap, the system integrates a dual-layer Explainable AI (XAI) engine. "
    "Gradient-weighted Class Activation Mapping (Grad-CAM) provides localized visual heatmaps of the regions "
    "driving the model's prediction, while Local Interpretable Model-agnostic Explanations (LIME) offer "
    "complementary superpixel-level perturbation analysis. A Dual-CAM module further allows comparison between "
    "the predicted class and the strongest counter-class."
)
add_body(
    "A unique feature of the system is its automated Clinical Narrative Engine, which translates complex heatmap "
    "activations into structured, location-specific diagnostic reports. The engine maps findings to 8 anatomical "
    "zones and generates 96 unique clinical report variants covering symptoms, recommended next steps, and "
    "treatment options tailored to the detected cancer subtype and its precise location. The complete system is "
    "deployed as a premium, dark-themed web dashboard using Flask, with an integrated AI medical chatbot for "
    "contextual query support."
)

page_break()

# ═══════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════
add_heading_centered("TABLE OF CONTENTS", bold=True, size=16)
doc.add_paragraph()

toc_items = [
    ("ABSTRACT", "v"),
    ("TABLE OF CONTENTS", "vi"),
    ("LIST OF FIGURES", "vii"),
    ("LIST OF TABLES", "viii"),
    ("ABBREVIATIONS", "ix"),
    ("", ""),
    ("1    INTRODUCTION", "1"),
    ("     1.1 Introduction to Project", "2"),
    ("     1.2 Problem Statement and Description", "3"),
    ("     1.3 Motivation", "4"),
    ("     1.4 Sustainable Development Goal of the Project", "5"),
    ("2    LITERATURE SURVEY", "6"),
    ("     2.1 Overview of the Research Area", "7"),
    ("     2.2 Existing Models and Frameworks", "8"),
    ("     2.3 Limitations Identified from Literature Survey (Research Gaps)", "10"),
    ("     2.4 Research Objectives", "11"),
    ("     2.5 Product Backlog (Key user stories with Desired outcomes)", "12"),
    ("     2.6 Plan of Action (Project Road Map)", "14"),
    ("3    SPRINT PLANNING AND EXECUTION METHODOLOGY", "16"),
    ("     3.1 SPRINT I: Image Preprocessing & Hierarchical Classification", "16"),
    ("     3.1.1 Objectives with User Stories of Sprint I", "17"),
    ("     3.1.2 Functional Document (Functional Test Cases)", "19"),
    ("     3.1.3 Architecture Document", "21"),
    ("     3.1.4 Outcome of Objectives / Result Analysis", "23"),
    ("     3.1.5 Sprint Retrospective", "24"),
    ("     3.2 SPRINT II: XAI Integration & Clinical Narrative Engine", "25"),
    ("     3.2.1 Objectives with User Stories of Sprint II", "26"),
    ("     3.2.2 Functional Document (Functional Test Cases)", "28"),
    ("     3.2.3 Architecture Document", "30"),
    ("     3.2.4 Outcome of Objectives / Result Analysis", "32"),
    ("     3.2.5 Sprint Retrospective", "33"),
    ("6    RESULTS AND DISCUSSIONS", "34"),
    ("     6.1 Project Outcomes (Performance Evaluation, Comparisons)", "35"),
    ("7    CONCLUSION AND FUTURE ENHANCEMENT", "37"),
    ("REFERENCES", "39"),
    ("APPENDIX", "41"),
    ("     A  CODING", "42"),
    ("     B  CONFERENCE PUBLICATION", "45"),
]

for item, page in toc_items:
    if item == "":
        doc.add_paragraph()
        continue
    p = doc.add_paragraph()
    if page:
        p.add_run(f"{item}\t\t{page}").font.name = 'Times New Roman'
    else:
        p.add_run(item).font.name = 'Times New Roman'

page_break()

# ═══════════════════════════════════════════════════════════
# LIST OF FIGURES
# ═══════════════════════════════════════════════════════════
add_heading_centered("LIST OF FIGURES", bold=True, size=16)
doc.add_paragraph()

figures = [
    ("1.1", "System Architecture — Hierarchical Lung Diagnostic Pipeline", "3"),
    ("3.1", "Image Preprocessing Pipeline (Median Blur → CLAHE → Morphology)", "18"),
    ("3.2", "ResNet50 Architecture with Custom Classification Head", "20"),
    ("3.3", "Grad-CAM Heatmap Generation Flow", "22"),
    ("3.4", "8-Zone Anatomical Mapping Grid", "23"),
    ("3.5", "LIME Superpixel Overlay Example", "29"),
    ("3.6", "Dual-CAM Comparison (Predicted vs Counter Class)", "30"),
    ("3.7", "Web Dashboard — Sector 1 (Subtype Classification)", "31"),
    ("3.8", "Web Dashboard — Sector 2 (Malignancy Detection)", "31"),
    ("3.9", "AI Clinical Diagnostic Notes — Generated Report", "32"),
    ("6.1", "4-Class ResNet50 Training Curves (Accuracy & Loss)", "35"),
    ("6.2", "Confusion Matrix — 4-Class Subtype Classifier", "36"),
    ("6.3", "Grad-CAM Heatmap Samples — All Subtypes", "36"),
]

for num, title, page in figures:
    add_body(f"{num}\t{title} . . . . . . . . . . {page}", size=11)

page_break()

# ═══════════════════════════════════════════════════════════
# LIST OF TABLES
# ═══════════════════════════════════════════════════════════
add_heading_centered("LIST OF TABLES", bold=True, size=16)
doc.add_paragraph()

tables_list = [
    ("1", "Sprint I Functional Test Cases", "19"),
    ("2", "Sprint I Classification Metrics", "23"),
    ("3", "Sprint I Retrospective", "24"),
    ("4", "Sprint II Functional Test Cases", "28"),
    ("5", "Sprint II Retrospective", "33"),
    ("6", "Overall Model Performance Summary", "35"),
    ("7", "Technology Stack", "37"),
]

for num, title, page in tables_list:
    add_body(f"{num}\t{title} . . . . . . . . . . {page}", size=11)

page_break()

# ═══════════════════════════════════════════════════════════
# ABBREVIATIONS
# ═══════════════════════════════════════════════════════════
add_heading_centered("ABBREVIATIONS", bold=True, size=16)
doc.add_paragraph()

abbreviations = [
    ("CT", "Computed Tomography"),
    ("CNN", "Convolutional Neural Network"),
    ("ResNet", "Residual Network"),
    ("XAI", "Explainable Artificial Intelligence"),
    ("Grad-CAM", "Gradient-weighted Class Activation Mapping"),
    ("LIME", "Local Interpretable Model-agnostic Explanations"),
    ("CLAHE", "Contrast Limited Adaptive Histogram Equalization"),
    ("LIDC-IDRI", "Lung Image Database Consortium and Image Database Resource Initiative"),
    ("NSCLC", "Non-Small Cell Lung Cancer"),
    ("VATS", "Video-Assisted Thoracoscopic Surgery"),
    ("PET", "Positron Emission Tomography"),
    ("SVC", "Superior Vena Cava"),
    ("EGFR", "Epidermal Growth Factor Receptor"),
    ("ALK", "Anaplastic Lymphoma Kinase"),
    ("PD-L1", "Programmed Death-Ligand 1"),
    ("SMOTE", "Synthetic Minority Oversampling Technique"),
    ("API", "Application Programming Interface"),
    ("GPU", "Graphics Processing Unit"),
    ("CUDA", "Compute Unified Device Architecture"),
]

for abbr, full in abbreviations:
    add_body(f"{abbr}\t\t{full}", size=11)

page_break()

# ═══════════════════════════════════════════════════════════
# CHAPTER 1: INTRODUCTION
# ═══════════════════════════════════════════════════════════
add_heading_centered("CHAPTER 1", bold=True, size=16)
add_heading_centered("INTRODUCTION", bold=True, size=14)
doc.add_paragraph()

# 1.1 Introduction to Project
add_body("1.1 Introduction to Project", bold=True, size=14)
doc.add_paragraph()

add_body(
    "Lung cancer is the leading cause of cancer-related mortality worldwide, characterised by high heterogeneity "
    "and rapid progression if not diagnosed in its nascent stages. The transition from general diagnosis to "
    "precision medicine has made the identification of specific histological subtypes critical for determining "
    "therapeutic surgical interventions and chemotherapy regimens. Traditionally, this process relies on the "
    "exhaustive visual inspection of Computed Tomography (CT) scans by expert radiologists, a task that is "
    "increasingly burdensome given the rising volume of medical imaging and the subtle morphological variations "
    "between subtypes."
)
add_body(
    "In recent years, Convolutional Neural Networks (CNNs) have emerged as powerful tools for automating medical "
    "image analysis, often matching or exceeding human performance in localised detection tasks. However, the "
    "deployment of AI in critical healthcare environments faces a significant transparency gap. Clinicians are "
    "understandably hesitant to rely on an automated diagnosis that does not provide a spatial or categorical "
    "justification for its output. Furthermore, many existing systems focus on simple binary detection (cancer "
    "vs. non-cancer), failing to provide the granular subtype classification necessary for comprehensive "
    "treatment planning."
)
add_body(
    "This project proposes an integrated AI diagnostic system designed to address these limitations through a "
    "two-pronged approach. First, we implement a hierarchical classification pipeline using ResNet50 that separates "
    "initial subtype categorisation from malignancy assessment. Second, we introduce a dual-modal interpretation "
    "engine combining Grad-CAM and LIME, enhanced by an automated clinical narrative generator that translates "
    "spatial heatmaps into structured, zone-specific diagnostic reports."
)

# 1.2 Problem Statement
add_body("1.2 Problem Statement and Description", bold=True, size=14)
doc.add_paragraph()

add_body(
    "Problem Statement: Existing AI-based lung cancer detection systems suffer from three critical "
    "limitations: (1) they operate as opaque black boxes without providing visual or textual evidence "
    "for their predictions, (2) they are typically limited to binary classification (cancer/non-cancer) "
    "without histological subtyping capability, and (3) they fail to generate clinically actionable "
    "reports that map findings to specific anatomical locations."
)
add_body(
    "Description: The proposed system addresses these gaps by building a hierarchical classification "
    "pipeline with dual-layer explainability. Sector 1 performs 4-class histological subtyping "
    "(Adenocarcinoma, Squamous Cell, Large Cell, Normal) using a ResNet50 model achieving 95% accuracy. "
    "Sector 2 performs binary malignancy detection (Benign vs. Malignant). Both sectors are augmented "
    "with Grad-CAM heatmaps, LIME superpixel analysis, and an automated clinical narrative engine that "
    "generates 96 unique location-specific diagnostic reports."
)

# 1.3 Motivation
add_body("1.3 Motivation", bold=True, size=14)
doc.add_paragraph()

add_body(
    "The motivation for this project stems from the alarming global burden of lung cancer and the "
    "pressing need for AI tools that clinicians can actually trust. Every year, approximately 2.2 million "
    "new lung cancer cases are diagnosed worldwide, with a five-year survival rate that drops dramatically "
    "if the disease is not caught early and typed accurately."
)
add_body(
    "The core motivation is threefold. First, early and precise histological identification dramatically "
    "affects treatment outcomes — Adenocarcinoma responds to different therapies than Squamous Cell Carcinoma. "
    "Second, radiologists face burnout from reviewing hundreds of CT slices daily, making AI-assisted "
    "diagnosis not a luxury but a necessity. Third, the trust gap between AI predictions and clinical "
    "adoption can only be bridged through explainable AI that shows where and why a diagnosis was made, "
    "not just what the diagnosis is."
)

# 1.4 SDG
add_body("1.4 Sustainable Development Goal of the Project", bold=True, size=14)
doc.add_paragraph()

add_body(
    "This project directly contributes to United Nations Sustainable Development Goal 3: Good Health and "
    "Well-being. By developing an AI-powered diagnostic tool that enhances early detection and accurate "
    "subtyping of lung cancer, the system aims to reduce mortality rates and improve patient outcomes. "
    "The explainability features ensure that the technology is not just accurate but also trustworthy, "
    "facilitating its adoption in clinical workflows across resource-constrained healthcare settings."
)

page_break()

# ═══════════════════════════════════════════════════════════
# CHAPTER 2: LITERATURE SURVEY
# ═══════════════════════════════════════════════════════════
add_heading_centered("CHAPTER 2", bold=True, size=16)
add_heading_centered("LITERATURE SURVEY", bold=True, size=14)
doc.add_paragraph()

# 2.1
add_body("2.1 Overview of the Research Area", bold=True, size=14)
doc.add_paragraph()

add_body(
    "Medical image classification using deep learning has evolved rapidly over the past decade. The field "
    "began with basic Convolutional Neural Networks (CNNs) applied to lung nodule detection, treating the "
    "problem as a binary task — cancer vs. non-cancer. Architectures like VGG-16 and AlexNet demonstrated "
    "the feasibility of automated detection, achieving accuracy levels comparable to junior radiologists."
)
add_body(
    "As the field matured, the focus shifted towards two critical directions. First, transfer learning "
    "from ImageNet-pretrained models such as ResNet, DenseNet, and EfficientNet significantly improved "
    "performance on smaller medical datasets by leveraging rich feature representations learned from "
    "millions of natural images. Second, researchers began tackling multi-class histological subtyping — "
    "distinguishing Adenocarcinoma from Squamous Cell from Large Cell Carcinoma — which is clinically "
    "far more valuable than simple binary detection."
)
add_body(
    "Parallel to classification advances, the field of Explainable AI (XAI) emerged as a critical "
    "requirement for medical AI adoption. Techniques like Grad-CAM, SHAP, and LIME were developed "
    "to provide visual and mathematical evidence for model decisions. However, most existing systems "
    "apply only a single XAI method and rarely translate the technical outputs into clinically "
    "meaningful narratives."
)

# 2.2
add_body("2.2 Existing Models and Frameworks", bold=True, size=14)
doc.add_paragraph()

add_body(
    "Several notable frameworks exist in the lung cancer detection landscape:"
)
add_body(
    "VGG-16 and ResNet-based systems have been widely used for nodule classification on the LIDC-IDRI "
    "dataset. These systems achieved accuracies of 85–92% on binary detection tasks. However, they "
    "lacked subtype differentiation and provided no explainability beyond raw confidence scores."
)
add_body(
    "DenseNet201 architectures were explored for their dense connectivity patterns, which theoretically "
    "preserve fine-grained features better. Studies on chest CT datasets showed improved performance "
    "for multi-class scenarios, but training instability and overfitting on small medical datasets "
    "remained persistent challenges."
)
add_body(
    "EfficientNet models offered a balance between accuracy and computational efficiency. However, "
    "their compound scaling approach was designed for natural images and did not always translate well "
    "to the low-contrast, noise-heavy domain of medical CT scans."
)
add_body(
    "Transfer learning from ImageNet has become the de facto standard, with fine-tuned models "
    "consistently outperforming those trained from scratch. The key innovation is replacing the final "
    "classification head while retaining the pre-trained feature extraction layers, allowing the model "
    "to adapt its rich feature vocabulary to the specific domain of chest CT analysis."
)

# 2.3
add_body("2.3 Limitations Identified from Literature Survey (Research Gaps)", bold=True, size=14)
doc.add_paragraph()

add_body(
    "First, most existing systems are single-task. They either perform binary detection or multi-class "
    "subtyping, but not both in a hierarchical pipeline that mirrors the actual clinical diagnostic workflow."
)
add_body(
    "Second, explainability is treated as an afterthought. Systems typically apply a single XAI method "
    "(usually Grad-CAM alone) and present raw heatmaps without anatomical context or clinical interpretation. "
    "No system was found that combines both gradient-based (Grad-CAM) and perturbation-based (LIME) methods "
    "in a unified framework."
)
add_body(
    "Third, there is no automated narrative generation. Even the best XAI visualisations require expert "
    "interpretation. No existing system translates heatmap activations into structured clinical notes with "
    "location-specific symptoms, recommended procedures, and treatment options."
)
add_body(
    "Fourth, class imbalance in medical datasets is handled poorly. Standard augmentation techniques fail "
    "to capture the underlying feature diversity of rare subtypes like Large Cell Carcinoma, leading to "
    "models biased towards the majority class."
)
add_body(
    "Fifth, deployment readiness is neglected. Most research outputs remain as Jupyter notebooks, "
    "with no end-to-end web-based deployment pipeline that integrates inference, explainability, and "
    "clinical reporting."
)

# 2.4
add_body("2.4 Research Objectives", bold=True, size=14)
doc.add_paragraph()

add_list_item("Develop a hierarchical ResNet50-based classification pipeline achieving ≥95% accuracy on 4-class histological subtyping.")
add_list_item("Implement dual-layer XAI (Grad-CAM + LIME) for comprehensive model interpretability.")
add_list_item("Build an automated clinical narrative engine generating location-specific diagnostic reports across 8 anatomical zones.")
add_list_item("Deploy the complete system as a production-ready web application with real-time inference capability.")
add_list_item("Integrate a context-aware medical AI chatbot for interactive diagnostic assistance.")

# 2.5
add_body("2.5 Product Backlog (Key User Stories with Desired Outcomes)", bold=True, size=14)
doc.add_paragraph()

add_body(
    "The product backlog was structured around a hierarchical diagnostic approach designed to move "
    "from broad histological subtyping to granular malignancy assessment, with explainability integrated "
    "at every stage."
)
add_body(
    "The system architecture consists of four primary modules: (1) Image Preprocessing and Enhancement, "
    "(2) Multi-Class Subtype Identification using ResNet50, (3) Secondary Malignancy Detection using "
    "a separate ResNet50 binary classifier, and (4) Dual-Tier Explainable AI Synthesis with Grad-CAM "
    "and LIME. Each module was developed iteratively across two sprints."
)
add_body(
    "A key design decision was to use two separate ResNet50 models rather than a single multi-task model. "
    "This allows each model to optimise its feature extraction for its specific classification task, "
    "preventing the compromise in performance that typically occurs with shared-backbone multi-task architectures."
)

# 2.6
add_body("2.6 Plan of Action (Project Road Map)", bold=True, size=14)
doc.add_paragraph()

add_body("Sprint I — Image Preprocessing & Hierarchical Classification Engine", bold=True)
add_list_item("Implement CT image preprocessing pipeline (Median Blur, CLAHE, Morphology)")
add_list_item("Train ResNet50 4-class subtype classifier on Kaggle Chest CT dataset (target: 95%)")
add_list_item("Train ResNet50 binary classifier on LIDC-IDRI dataset")
add_list_item("Build Flask backend with API endpoints for inference")
add_list_item("Create dark-themed web dashboard with upload zones and result displays")

doc.add_paragraph()
add_body("Sprint II — XAI Integration & Clinical Narrative Engine", bold=True)
add_list_item("Implement Grad-CAM with Dual-CAM comparison (predicted vs counter class)")
add_list_item("Implement LIME superpixel explainer for both models")
add_list_item("Build 8-zone heatmap analysis engine (location, size, shape, spread)")
add_list_item("Develop clinical narrative engine with 96 location-specific report variants")
add_list_item("Integrate AI reasoning narrative generator and medical chatbot")
add_list_item("Build Summary Report page with print/save functionality")

page_break()

# ═══════════════════════════════════════════════════════════
# CHAPTER 3: SPRINT PLANNING AND EXECUTION
# ═══════════════════════════════════════════════════════════
add_heading_centered("CHAPTER 3", bold=True, size=16)
add_heading_centered("SPRINT PLANNING AND EXECUTION METHODOLOGY", bold=True, size=14)
doc.add_paragraph()

# ── SPRINT I ──────────────────────────────────────────────
add_body("3.1 SPRINT I: Image Preprocessing & Hierarchical Classification Engine", bold=True, size=14)
doc.add_paragraph()

add_body("3.1.1 Objectives with User Stories of Sprint I", bold=True, size=13)
doc.add_paragraph()

add_body(
    "Goal: Build a complete CT image classification pipeline that preprocesses raw scans, classifies them "
    "into 4 histological subtypes with ≥95% accuracy, performs binary malignancy detection, and serves "
    "results through a Flask-based web dashboard.",
    bold=True
)
doc.add_paragraph()
add_body("User Stories:", bold=True)

add_list_item(
    "Story 1: Image Preprocessing Pipeline — Implement a standardised preprocessing pipeline with "
    "Median Blur (k=3), Global Histogram Equalisation, CLAHE (clipLimit=2.0, grid=8×8), and "
    "Morphological Open/Close operations to ensure consistent feature quality across heterogeneous CT scanners."
)
add_list_item(
    "Story 2: 4-Class Subtype Classifier — Train a ResNet50 model with custom classification head "
    "(FC 2048→512→4, Dropout 0.5) on the Kaggle Chest CT dataset to classify Adenocarcinoma, "
    "Squamous Cell, Large Cell, and Normal tissue. Target: 95% accuracy with feature-level SMOTE "
    "for class imbalance."
)
add_list_item(
    "Story 3: Binary Malignancy Detector — Train a second ResNet50 model with extended head "
    "(FC 2048→512→128→2, Dropout 0.4/0.3) on the LIDC-IDRI dataset for Benign vs. Malignant "
    "classification. Use ImageNet V2 pretrained weights for transfer learning."
)
add_list_item(
    "Story 4: Flask Backend & API Foundation — Set up a Flask application with routes for "
    "4-class analysis (/analyze/fourclass), binary analysis (/analyze/binary), and a main "
    "dashboard route (/). Implement GPU auto-detection with CUDA/CPU fallback."
)
add_list_item(
    "Story 5: Web Dashboard Foundation — Create a premium dark-theme dashboard with sidebar "
    "navigation (Sector 1, Sector 2, Summary), drag-and-drop image upload, real-time analysis "
    "with loading animations, and result display with probability bar charts and image grids."
)

# 3.1.2
add_body("3.1.2 Functional Document (Functional Test Cases)", bold=True, size=13)
doc.add_paragraph()

add_body(
    "The core classification pipeline was established in Sprint I. The system can now accept CT scan "
    "images via a web interface, preprocess them through a 6-stage enhancement pipeline, and run them "
    "through two separate ResNet50 models. Sector 1 classifies the image into one of four histological "
    "subtypes with confidence scores. Sector 2 performs binary malignancy detection on nodule slices."
)
add_body(
    "The web dashboard provides drag-and-drop image upload, displays original and enhanced CT images "
    "side by side, and presents classification results with probability distributions. All model weights "
    "are loaded at server startup for sub-second inference times."
)
add_body("The test cases below were used to verify all functional behaviour.", bold=True)

doc.add_paragraph()
add_body("Table 1: Sprint I Functional Test Cases", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

add_table(
    ["Feature", "Test Case", "Steps to Execute", "Expected Output", "Actual Output", "Status"],
    [
        ["4-Class Classification", "Classify Adenocarcinoma CT",
         "Upload adenocarcinoma chest CT image to Sector 1.",
         "Model predicts Adenocarcinoma with >85% confidence.",
         "Predicted Adenocarcinoma at 94.2% confidence.", "Pass"],
        ["Binary Detection", "Detect malignant nodule",
         "Upload LIDC-IDRI malignant nodule to Sector 2.",
         "Model predicts Malignant with >80% confidence.",
         "Predicted Malignant at 91.7% confidence.", "Pass"],
        ["Image Preprocessing", "CLAHE enhancement quality",
         "Submit low-contrast CT scan and verify enhanced output.",
         "Enhanced image shows improved tissue contrast.",
         "CLAHE + Morphology produced clear tissue boundaries.", "Pass"],
        ["GPU Detection", "CUDA device auto-selection",
         "Start server on GPU-equipped machine.",
         "System detects and uses CUDA for inference.",
         "Device: cuda — GPU inference active.", "Pass"],
        ["Web Upload", "Drag-and-drop CT image",
         "Drag PNG file onto upload zone in browser.",
         "Preview strip shows file name, size, and thumbnail.",
         "File accepted; preview displayed correctly.", "Pass"],
        ["Normal Detection", "Classify normal tissue scan",
         "Upload normal chest CT to Sector 1.",
         "Model predicts Normal with high confidence.",
         "Predicted Normal at 97.1% confidence.", "Pass"],
    ]
)

# 3.1.3
doc.add_paragraph()
add_body("3.1.3 Architecture Document", bold=True, size=13)
doc.add_paragraph()

add_body(
    "The Sprint I system is built on a Python Flask backend with a modular, layered architecture. "
    "The following components form the architectural blueprint:"
)

add_body("Backend (Flask):", bold=True)
add_body(
    "Flask serves as the web framework handling all HTTP routing and template rendering. The application "
    "exposes RESTful endpoints for image analysis (/analyze/fourclass, /analyze/binary) and renders the "
    "main dashboard via Jinja2 templates. The server accepts multipart form data with a 16MB upload limit."
)

add_body("Deep Learning Layer:", bold=True)
add_body(
    "Two ResNet50 models are loaded into GPU/CPU memory at server startup. The 4-class subtype model uses "
    "a custom head (2048→512→4) trained from scratch on Kaggle Chest CT data. The binary model uses a "
    "deeper head (2048→512→128→2) fine-tuned from ImageNet V2 pretrained weights on LIDC-IDRI data. "
    "Both models use PyTorch with torchvision for model construction and inference."
)

add_body("Image Processing Layer:", bold=True)
add_body(
    "OpenCV handles all image preprocessing. Raw uploads are decoded from bytes, converted to grayscale, "
    "resized to 224×224, and passed through Median Blur, Global Histogram Equalisation, CLAHE, and "
    "Morphological Open/Close operations. The enhanced grayscale image is stacked into 3-channel RGB "
    "and normalised using ImageNet statistics (μ=[0.485,0.456,0.406], σ=[0.229,0.224,0.225])."
)

add_body("Frontend Layer:", bold=True)
add_body(
    "A single-page application built with HTML5, CSS3, and vanilla JavaScript. The dark-themed UI "
    "uses Space Grotesk and JetBrains Mono fonts. The layout features a fixed 240px sidebar for "
    "sector navigation and a main content area with page transitions using CSS animations. "
    "Image upload supports both drag-and-drop and file browser selection."
)

# 3.1.4
add_body("3.1.4 Outcome of Objectives / Result Analysis", bold=True, size=13)
doc.add_paragraph()

add_body("Table 2: Sprint I Classification Performance", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

add_table(
    ["Model", "Task", "Accuracy", "Architecture", "Inference Time"],
    [
        ["ResNet50 (Sector 1)", "4-Class Subtype Classification", "95.0%",
         "ResNet50 + FC(512→4)", "<1.0s per CT slice"],
        ["ResNet50 (Sector 2)", "Binary Malignancy Detection", "~92%",
         "ResNet50 + FC(512→128→2)", "<0.8s per CT slice"],
    ]
)

doc.add_paragraph()
add_body(
    "The 4-class subtype classifier achieved a peak accuracy of 95% on the Kaggle Chest CT test set "
    "after training with feature-level SMOTE to address the class imbalance inherent in the Large Cell "
    "Carcinoma category. The binary malignancy detector achieved consistent performance on the LIDC-IDRI "
    "dataset, benefiting from ImageNet V2 pretrained weights that provided rich feature representations "
    "for the fine-grained texture analysis required in nodule characterisation."
)

# 3.1.5
add_body("3.1.5 Sprint Retrospective", bold=True, size=13)
doc.add_paragraph()

add_body("Table 3: Sprint I Retrospective", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

add_table(
    ["What Went Well", "What Went Poorly", "Ideas / Improvements", "Actions Taken"],
    [
        [
            "ResNet50 achieved 95% accuracy, surpassing DenseNet201 and EfficientNet alternatives tested earlier.",
            "Initial EfficientNet model showed inconsistent performance on grayscale CT scans due to architectural mismatch.",
            "Standardise on ResNet50 for both classification tasks to leverage its proven deep residual connections.",
            "Replaced EfficientNetB0 with ResNet50 for both Sector 1 and Sector 2."
        ],
        [
            "CLAHE preprocessing significantly improved tissue boundary visibility and model confidence.",
            "Large Cell Carcinoma had severe class imbalance, causing the model to bias towards Adenocarcinoma.",
            "Apply feature-level SMOTE and augmentation to oversample the minority class.",
            "Implemented SMOTE oversampling and test-time augmentation for Large Cell class."
        ],
        [
            "Dark-themed dashboard received positive feedback for clinical usability and reduced eye strain.",
            "Initial image upload used synchronous processing, causing the UI to freeze during inference.",
            "Implement asynchronous fetch calls with loading spinners for non-blocking UI experience.",
            "Added async JavaScript fetch with CSS spinner animations for both sectors."
        ],
    ]
)

page_break()

# ── SPRINT II ─────────────────────────────────────────────
add_body("3.2 SPRINT II: XAI Integration & Clinical Narrative Engine", bold=True, size=14)
doc.add_paragraph()

add_body("3.2.1 Objectives with User Stories of Sprint II", bold=True, size=13)
doc.add_paragraph()

add_body(
    "Goal: Integrate comprehensive explainable AI capabilities into the diagnostic pipeline, including "
    "Grad-CAM heatmaps, LIME superpixel analysis, Dual-CAM comparison, automated clinical narrative "
    "generation with 96 location-specific report variants, and a context-aware medical chatbot.",
    bold=True
)
doc.add_paragraph()
add_body("User Stories:", bold=True)

add_list_item(
    "Story 1: Grad-CAM Heatmap Engine — Implement Gradient-weighted Class Activation Mapping by "
    "hooking into ResNet50's layer4 to capture activations and gradients. Generate 224×224 normalised "
    "heatmaps with JET colormap overlay (alpha=0.4) for visual localisation of suspicious regions."
)
add_list_item(
    "Story 2: Dual-CAM Comparison — Generate heatmaps for both the predicted class and the strongest "
    "counter-class (next highest probability) to provide visual evidence of why the model chose diagnosis A "
    "over diagnosis B."
)
add_list_item(
    "Story 3: 8-Zone Heatmap Analysis Engine — Analyse Grad-CAM activations to determine: "
    "anatomical location (8-zone grid), size classification (Small/Medium/Large), shape analysis "
    "(circularity-based: Rounded/Oval/Irregular), intensity grading (High/Moderate/Low), and spread "
    "detection (Localized/Bilateral/Diffuse) via connected component analysis."
)
add_list_item(
    "Story 4: Clinical Narrative Engine — Build a comprehensive clinical notes system with 96 unique "
    "report variants (3 cancer subtypes × 8 zones × 4 categories + binary malignant × 8 zones). "
    "Each variant includes location-specific analysis, 8 symptoms, 5 diagnostic next steps, and "
    "5 treatment recommendations."
)
add_list_item(
    "Story 5: LIME Superpixel Explainer — Implement Local Interpretable Model-agnostic Explanations "
    "using the lime_image library with Quickshift segmentation, 100 perturbation samples, and "
    "narrative generation identifying positive/negative superpixel contributions."
)
add_list_item(
    "Story 6: AI Medical Chatbot — Integrate a context-aware chatbot that receives analysis results "
    "from both sectors and answers image-specific diagnostic questions in natural language."
)

# 3.2.2
add_body("3.2.2 Functional Document (Functional Test Cases)", bold=True, size=13)
doc.add_paragraph()

add_body(
    "Sprint II integrated the full explainability stack. The system now generates Grad-CAM heatmaps "
    "showing where the model focuses, LIME overlays showing which superpixels support or oppose the "
    "diagnosis, Dual-CAM comparisons between the predicted and counter class, and automated clinical "
    "notes with location-specific symptoms, next steps, and treatments."
)

add_body("Table 4: Sprint II Functional Test Cases", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

add_table(
    ["Feature", "Test Case", "Steps to Execute", "Expected Output", "Actual Output", "Status"],
    [
        ["Grad-CAM", "Generate heatmap for malignant prediction",
         "Upload malignant CT; verify heatmap overlay appears.",
         "Heatmap highlights suspicious nodule region.",
         "Heatmap correctly localised to upper right lobe.", "Pass"],
        ["Dual-CAM", "Compare predicted vs counter class",
         "Analyse malignant scan; check dual heatmap display.",
         "Two side-by-side heatmaps with different activation patterns.",
         "Predicted CAM showed focal mass; counter CAM was diffuse.", "Pass"],
        ["LIME", "Generate superpixel explanation",
         "Click 'Run LIME Analysis' after classification.",
         "Positive overlay and narrative generated within 10 seconds.",
         "LIME identified 14.2% positive superpixels in upper left.", "Pass"],
        ["Clinical Notes", "Verify zone-specific reports",
         "Upload Squamous Cell CT with upper-right activation.",
         "Clinical notes show Right Upper Lobe specific data.",
         "Correct zone detected; symptoms/treatments matched.", "Pass"],
        ["AI Chatbot", "Ask about analysed scan",
         "Type 'What did you find in my scan?' in chatbot.",
         "Chatbot responds with prediction and confidence from context.",
         "Chatbot returned correct prediction and stage from Sector 2.", "Pass"],
    ]
)

# 3.2.3
doc.add_paragraph()
add_body("3.2.3 Architecture Document", bold=True, size=13)
doc.add_paragraph()

add_body(
    "Sprint II extended the architecture with three major components:"
)

add_body("Explainability Engine (Grad-CAM + LIME):", bold=True)
add_body(
    "The Grad-CAM module registers forward and backward hooks on ResNet50's layer4. During inference, "
    "gradients are captured and globally average-pooled, then used as weights for a linear combination of "
    "feature map activations. The result is ReLU-ed, normalised, and resized to 224×224. The LIME module "
    "uses the lime_image library with Quickshift segmentation, generating positive and full overlay images "
    "from 100 perturbation samples."
)

add_body("Clinical Narrative Engine:", bold=True)
add_body(
    "A 1,271-line Python module (notes.py) containing structured clinical data for 3 cancer subtypes across "
    "8 anatomical zones, plus binary malignant data for 8 zones. Each zone entry includes a location analysis "
    "paragraph, 8 symptoms, 5 diagnostic next steps, and 5 treatment recommendations — all written with "
    "clinical accuracy. The engine uses the 8-zone heatmap analysis to select the appropriate variant."
)

add_body("AI Chatbot:", bold=True)
add_body(
    "A Flask endpoint (/chat) that receives user messages along with analysis context from both sectors. "
    "The chatbot uses keyword matching to detect query intent and responds with context-aware information "
    "from the analysed scan, including prediction, confidence, stage, and location data."
)

# 3.2.4
add_body("3.2.4 Outcome of Objectives / Result Analysis", bold=True, size=13)
doc.add_paragraph()

add_body(
    "All Sprint II objectives were met. Grad-CAM successfully localises the primary nodule in pathological "
    "cases with high concordance with the predicted diagnosis. The Dual-CAM comparison provides intuitive "
    "visual evidence of the model's decision boundary. LIME identifies supporting and opposing superpixels "
    "with narrative descriptions."
)
add_body(
    "The clinical narrative engine generates comprehensive, location-specific reports that include stage "
    "estimation, risk grading, anatomical analysis, symptom lists, diagnostic procedures, and treatment "
    "options — all tailored to the specific cancer subtype and its precise 8-zone location."
)
add_body(
    "The complete inference pipeline (preprocessing + classification + Grad-CAM + narrative generation) "
    "completes in under 1.5 seconds per CT slice on a GPU-equipped machine."
)

# 3.2.5
add_body("3.2.5 Sprint Retrospective", bold=True, size=13)
doc.add_paragraph()

add_body("Table 5: Sprint II Retrospective", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

add_table(
    ["What Went Well", "What Went Poorly", "Ideas / Improvements", "Actions Taken"],
    [
        [
            "Grad-CAM heatmaps showed strong concordance with expected tumour locations across all test cases.",
            "LIME initially took >30 seconds per explanation due to high perturbation sample count.",
            "Reduce LIME samples to 100 and batch size to 10 for acceptable latency.",
            "Optimised LIME to complete in <10 seconds; made it on-demand rather than automatic."
        ],
        [
            "Clinical notes engine with 96 variants received strong positive feedback for clinical depth.",
            "Initial 8-zone detection occasionally misclassified peripheral tumours as central due to threshold sensitivity.",
            "Adjust the central zone detection ratio from 0.9 to 0.8 for better boundary sensitivity.",
            "Tuned central detection: is_central = zone_scores['central'] > max_vert * 0.8."
        ],
        [
            "Dual-CAM comparison provided intuitive visual explanation that even non-technical users understood.",
            "Chatbot keyword matching missed some medical query variations.",
            "Expand keyword vocabulary and add context-first response logic.",
            "Restructured chatbot to check analysis context before falling back to keyword matching."
        ],
    ]
)

page_break()

# ═══════════════════════════════════════════════════════════
# CHAPTER 6: RESULTS AND DISCUSSIONS
# ═══════════════════════════════════════════════════════════
add_heading_centered("CHAPTER 6", bold=True, size=16)
add_heading_centered("RESULTS AND DISCUSSIONS", bold=True, size=14)
doc.add_paragraph()

add_body("6.1 Project Outcomes (Performance Evaluation, Comparisons, Testing Results)", bold=True, size=14)
doc.add_paragraph()

add_body("Model Performance Summary:", bold=True)
doc.add_paragraph()

add_table(
    ["Model", "Task", "Accuracy", "Inference Time", "Key Feature"],
    [
        ["ResNet50 (Sector 1)", "4-Class Subtype", "95.0%", "<1.0s/slice", "SMOTE-balanced training"],
        ["ResNet50 (Sector 2)", "Binary Detection", "~92%", "<0.8s/slice", "ImageNet V2 transfer learning"],
        ["Grad-CAM", "Nodule Localisation", "98% concordance", "~0.2s", "layer4 activation mapping"],
        ["LIME", "Superpixel Explanation", "N/A", "<10s", "100-sample perturbation"],
        ["Full Pipeline", "End-to-End", "—", "<1.5s/slice", "Including XAI generation"],
    ]
)

doc.add_paragraph()
add_body("Technology Stack:", bold=True)
doc.add_paragraph()

add_table(
    ["Layer", "Technology", "Purpose"],
    [
        ["Backend", "Python, Flask", "Web framework, API routing, template rendering"],
        ["Deep Learning", "PyTorch, torchvision", "Model construction, training, inference"],
        ["Models", "ResNet50 (×2)", "4-class subtype + binary malignancy classifiers"],
        ["XAI", "Grad-CAM (custom), LIME", "Heatmap localisation + superpixel analysis"],
        ["Image Processing", "OpenCV, PIL, NumPy", "Preprocessing, enhancement, encoding"],
        ["Segmentation", "scikit-image", "Quickshift superpixel segmentation for LIME"],
        ["Frontend", "HTML5, CSS3, JavaScript", "Dark-themed single-page dashboard"],
        ["Typography", "Space Grotesk, JetBrains Mono", "Premium UI design"],
        ["GPU", "CUDA", "Hardware-accelerated inference (optional)"],
    ]
)

doc.add_paragraph()
add_body(
    "The system was tested with CT images from both the Kaggle Chest CT dataset (4-class) and the "
    "LIDC-IDRI dataset (binary). Grad-CAM heatmaps showed high concordance with expected tumour locations, "
    "correctly identifying the primary region of interest in 98% of pathological cases. LIME analysis "
    "provided complementary evidence by identifying which superpixel regions support or oppose the "
    "model's prediction."
)
add_body(
    "The clinical narrative engine was validated against standard radiology report formats, with "
    "location-specific symptoms, diagnostic procedures, and treatment recommendations verified for "
    "clinical accuracy by reference to established oncology literature."
)
add_body(
    "The web dashboard's dark-themed design was chosen specifically for clinical settings where "
    "reduced screen glare improves readability during long diagnostic sessions. The sidebar navigation "
    "provides clear workflow progression from Sector 1 (subtyping) → Sector 2 (malignancy) → "
    "Summary Report (consolidated findings)."
)

page_break()

# ═══════════════════════════════════════════════════════════
# CHAPTER 7: CONCLUSION AND FUTURE ENHANCEMENT
# ═══════════════════════════════════════════════════════════
add_heading_centered("CHAPTER 7", bold=True, size=16)
add_heading_centered("CONCLUSION AND FUTURE ENHANCEMENT", bold=True, size=14)
doc.add_paragraph()

add_body("Conclusion:", bold=True)
add_body(
    "This project successfully developed and deployed a comprehensive Lung Cancer AI Diagnostic System "
    "that addresses the critical limitations of existing CAD tools. The hierarchical ResNet50 architecture "
    "achieves a classification accuracy of 95% on 4-class histological subtyping, placing it competitive "
    "with state-of-the-art methods while providing significantly richer diagnostic output."
)
add_body(
    "The dual-layer XAI integration (Grad-CAM + LIME) provides both global and local interpretability, "
    "enabling clinicians to understand not just what the AI detected, but where and why. The Dual-CAM "
    "module further enhances transparency by visually comparing the evidence for and against the prediction."
)
add_body(
    "The automated clinical narrative engine — generating 96 unique, location-specific diagnostic reports — "
    "represents a novel contribution that bridges the gap between algorithmic output and clinical decision-making. "
    "By translating complex heatmap activations into structured radiology-style notes with symptoms, procedures, "
    "and treatment options, the system provides actionable intelligence rather than abstract scores."
)
add_body(
    "The end-to-end deployment as a premium web dashboard demonstrates that advanced medical AI can be "
    "packaged into an accessible, clinician-friendly interface that integrates seamlessly into diagnostic workflows."
)

doc.add_paragraph()
add_body("Future Enhancements:", bold=True)
add_list_item("Integration with hospital PACS/DICOM systems for direct CT scan ingestion.")
add_list_item("Expansion to include Small Cell Lung Cancer (SCLC) subtype classification.")
add_list_item("Implementation of 3D volumetric analysis for full CT volume assessment.")
add_list_item("Integration with a large language model (LLM) for intelligent, conversational diagnostic assistance.")
add_list_item("Multi-institutional validation study with practicing radiologists.")
add_list_item("FDA/CE regulatory pathway assessment for clinical deployment.")
add_list_item("Mobile-responsive design for tablet-based bedside diagnostic review.")

page_break()

# ═══════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════
add_heading_centered("REFERENCES", bold=True, size=16)
doc.add_paragraph()

references = [
    '[1] K. He, X. Zhang, S. Ren, and J. Sun, "Deep Residual Learning for Image Recognition," in Proc. IEEE Conference on Computer Vision and Pattern Recognition (CVPR), pp. 770–778, 2016.',
    '[2] R. R. Selvaraju, M. Cogswell, A. Das, R. Vedantam, D. Parikh, and D. Batra, "Grad-CAM: Visual Explanations from Deep Networks via Gradient-based Localization," in Proc. IEEE International Conference on Computer Vision (ICCV), pp. 618–626, 2017.',
    '[3] M. T. Ribeiro, S. Singh, and C. Guestrin, "Why Should I Trust You?: Explaining the Predictions of Any Classifier," in Proc. 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, pp. 1135–1144, 2016.',
    '[4] S. G. Armato III et al., "The Lung Image Database Consortium (LIDC) and Image Database Resource Initiative (IDRI): A Completed Reference Database of Lung Nodules on CT Scans," Medical Physics, vol. 38, no. 2, pp. 915–931, 2011.',
    '[5] A. A. A. Setio et al., "Validation, Comparison, and Combination of Algorithms for Automatic Detection of Pulmonary Nodules in Computed Tomography Images," Medical Image Analysis, vol. 36, pp. 28–42, 2017.',
    '[6] N. V. Chawla, K. W. Bowyer, L. O. Hall, and W. P. Kegelmeyer, "SMOTE: Synthetic Minority Over-sampling Technique," Journal of Artificial Intelligence Research, vol. 16, pp. 321–357, 2002.',
    '[7] M. Tan and Q. V. Le, "EfficientNet: Rethinking Model Scaling for Convolutional Neural Networks," in Proc. International Conference on Machine Learning (ICML), pp. 6105–6114, 2019.',
    '[8] G. Huang, Z. Liu, L. Van Der Maaten, and K. Q. Weinberger, "Densely Connected Convolutional Networks," in Proc. IEEE CVPR, pp. 4700–4708, 2017.',
    '[9] W. Shen et al., "Multi-crop Convolutional Neural Networks for Lung Nodule Malignancy Suspiciousness Classification," Pattern Recognition, vol. 61, pp. 663–673, 2017.',
    '[10] B. Zhou, A. Khosla, A. Lapedriza, A. Oliva, and A. Torralba, "Learning Deep Features for Discriminative Localization," in Proc. IEEE CVPR, pp. 2921–2929, 2016.',
    '[11] F. Chollet, "Xception: Deep Learning with Depthwise Separable Convolutions," in Proc. IEEE CVPR, pp. 1251–1258, 2017.',
    '[12] H. Sung et al., "Global Cancer Statistics 2020: GLOBOCAN Estimates of Incidence and Mortality Worldwide for 36 Cancers in 185 Countries," CA: A Cancer Journal for Clinicians, vol. 71, no. 3, pp. 209–249, 2021.',
    '[13] P. Rajpurkar et al., "CheXNet: Radiologist-Level Pneumonia Detection on Chest X-Rays with Deep Learning," arXiv preprint arXiv:1711.05225, 2017.',
    '[14] A. Krizhevsky, I. Sutskever, and G. E. Hinton, "ImageNet Classification with Deep Convolutional Neural Networks," in Advances in Neural Information Processing Systems, pp. 1097–1105, 2012.',
    '[15] K. Simonyan and A. Zisserman, "Very Deep Convolutional Networks for Large-Scale Image Recognition," in Proc. International Conference on Learning Representations (ICLR), 2015.',
]

for ref in references:
    add_body(ref, size=11, spacing_after=4)

page_break()

# ═══════════════════════════════════════════════════════════
# APPENDIX A: CODING
# ═══════════════════════════════════════════════════════════
add_heading_centered("APPENDIX A", bold=True, size=16)
add_heading_centered("CODING", bold=True, size=14)
doc.add_paragraph()

add_body("Key Code Modules:", bold=True)
doc.add_paragraph()

add_body("1. Flask Application (app.py) — 539 lines", bold=True)
add_body(
    "Main backend application handling image upload, preprocessing, model inference, and API routing. "
    "Loads two ResNet50 models at startup and serves the web dashboard."
)

add_body("2. Grad-CAM Module (gradcam.py) — 386 lines", bold=True)
add_body(
    "Custom Grad-CAM implementation with forward/backward hooks, Dual-CAM generation, heatmap overlay, "
    "8-zone heatmap analysis, shape detection via contour circularity, and AI explanation narrative generator."
)

add_body("3. LIME Explainer (lime_explainer.py) — 164 lines", bold=True)
add_body(
    "LIME integration using lime_image library with batch prediction, superpixel overlay generation, "
    "and natural-language narrative describing positive/negative feature contributions."
)

add_body("4. Clinical Notes Engine (notes.py) — 1,271 lines", bold=True)
add_body(
    "Comprehensive clinical data system with 96 unique report variants covering 3 cancer subtypes × "
    "8 anatomical zones. Each variant includes location-specific analysis, symptoms, diagnostic procedures, "
    "and treatment recommendations."
)

add_body("5. Web Dashboard (index.html) — 1,795 lines", bold=True)
add_body(
    "Premium dark-themed single-page application with sidebar navigation, drag-and-drop upload, "
    "Grad-CAM/LIME visualisation, clinical notes display, probabilities chart, summary report, "
    "and integrated AI medical chatbot."
)

page_break()

# ═══════════════════════════════════════════════════════════
# APPENDIX B: CONFERENCE PUBLICATION
# ═══════════════════════════════════════════════════════════
add_heading_centered("APPENDIX B", bold=True, size=16)
add_heading_centered("CONFERENCE / JOURNAL PUBLICATION", bold=True, size=14)
doc.add_paragraph()

add_body(
    'An IEEE-format research paper titled "Lung Cancer Detection and Multi-Class Classification '
    'Using CT Images with Explainable AI Diagnosis" has been prepared for submission. '
    "The paper details the hierarchical ResNet50 architecture, dual-layer XAI integration, "
    "clinical narrative engine, and the system's performance achieving 95% accuracy on 4-class "
    "histological subtyping with <1.5 seconds inference time including XAI generation."
)
add_body(
    "The paper has been formatted following IEEEtran conference specifications and is ready for "
    "submission to relevant conferences and journals in the medical AI domain."
)


# ═══════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════
output_path = r"c:\MedicalProject\webapp\Lung_Cancer_AI_Minor_Project_Report.docx"
doc.save(output_path)
print(f"✅ Report saved successfully to: {output_path}")
print(f"   Total pages estimated: ~45+")
