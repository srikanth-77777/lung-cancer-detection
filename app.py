"""
================================================================
Flask Backend — Lung Cancer AI Diagnostic System
================================================================
Run: python app.py
Open: http://localhost:5000
================================================================
"""

import os
import cv2
import base64
import numpy as np
import torch
import torch.nn as nn
from flask import Flask, request, jsonify, render_template, send_file
from torchvision import models, transforms
from torchvision.models import ResNet50_Weights
from PIL import Image
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

from gradcam import (GradCAM, generate_heatmap_overlay,
                     analyze_heatmap,
                     generate_explanation_narrative)
from notes import generate_clinical_notes, generate_binary_clinical_notes
from lime_explainer import LimeExplainer, img_to_base64_rgb

from flask_sqlalchemy import SQLAlchemy
from datetime import datetime
import json

# ── Flask app ─────────────────────────────────────────────────
app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///webapp.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)

# ── Database Model ────────────────────────────────────────────
class Diagnosis(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)
    patient_name = db.Column(db.String(100), default="Anonymous")
    sector = db.Column(db.String(50))  # "Sector 1 (Binary)" or "Sector 2 (Subtype)"
    prediction = db.Column(db.String(100))
    confidence = db.Column(db.Float)
    is_malignant = db.Column(db.Boolean, nullable=True)
    stage = db.Column(db.String(50), nullable=True)
    location = db.Column(db.String(150), nullable=True)
    risk_level = db.Column(db.String(50), nullable=True)
    activation_pct = db.Column(db.Float, nullable=True)
    num_spots = db.Column(db.Integer, nullable=True)
    shape = db.Column(db.String(100), nullable=True)
    clinical_notes_json = db.Column(db.Text, nullable=True)  # Stringified JSON
    downloaded = db.Column(db.Boolean, default=False)

with app.app_context():
    db.create_all()

# ── Paths — load from environment or fallback ──────────────────
EFF_WEIGHTS     = os.getenv("EFF_WEIGHTS", r"C:\MedicalProject\best_ResNet50_v5.pth")
DENSE_WEIGHTS   = os.getenv("DENSE_WEIGHTS", r"C:\MedicalProject\chest_ct\resnet50_chest_ct_95acc_final.pth")
IMG_SIZE        = 224

# ── GPU ───────────────────────────────────────────────────────
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
print(f"Device: {device}")

# ── Class names ───────────────────────────────────────────────
BINARY_CLASSES = ['Benign', 'Malignant']
FOUR_CLASSES   = ['Adenocarcinoma', 'Large Cell',
                   'Normal', 'Squamous Cell']


# ────────────────────────────────────────────────────────────────
# MODEL DEFINITIONS
# ────────────────────────────────────────────────────────────────
class EnhancedCNN(nn.Module):
    def __init__(self, num_classes=2):
        super(EnhancedCNN, self).__init__()
        self.conv1 = nn.Sequential(
            nn.Conv2d(3, 64, 3, padding=1),
            nn.BatchNorm2d(64), nn.ReLU(), nn.MaxPool2d(2, 2))
        self.conv2 = nn.Sequential(
            nn.Conv2d(64, 32, 3, padding=1),
            nn.BatchNorm2d(32), nn.ReLU(), nn.MaxPool2d(2, 2))
        self.conv3 = nn.Sequential(
            nn.Conv2d(32, 32, 3, padding=1),
            nn.BatchNorm2d(32), nn.ReLU(), nn.Dropout2d(0.4))
        self.drop         = nn.Dropout(0.4)
        self.flatten_size = 32 * 56 * 56
        self.dense1 = nn.Sequential(
            nn.Linear(self.flatten_size, 256),
            nn.BatchNorm1d(256), nn.ReLU(), nn.Dropout(0.5))
        self.dense2 = nn.Sequential(
            nn.Linear(256, 128),
            nn.BatchNorm1d(128), nn.ReLU(), nn.Dropout(0.3))
        self.output = nn.Linear(128, num_classes)

    def forward(self, x):
        x = self.conv1(x); x = self.conv2(x)
        x = self.conv3(x); x = self.drop(x)
        x = x.view(x.size(0), -1)
        x = self.dense1(x); x = self.dense2(x)
        return self.output(x)


def build_efficientnet(num_classes=2):
    model = models.efficientnet_b0(
        weights=None)
    in_features = model.classifier[1].in_features
    model.classifier = nn.Sequential(
        nn.Dropout(0.4),
        nn.Linear(in_features, 256),
        nn.BatchNorm1d(256), nn.ReLU(),
        nn.Dropout(0.3),
        nn.Linear(256, 128), nn.ReLU(),
        nn.Linear(128, num_classes)
    )
    return model


def build_resnet50():
    m = models.resnet50(
        weights=ResNet50_Weights.IMAGENET1K_V2)
    m.fc = nn.Sequential(
        nn.Linear(m.fc.in_features, 512),
        nn.ReLU(),
        nn.Dropout(0.4),
        nn.Linear(512, 128),
        nn.ReLU(),
        nn.Dropout(0.3),
        nn.Linear(128, 2)
    )
    return m


def build_resnet50_subtype(num_classes=4):
    m = models.resnet50(weights=None)
    num_ftrs = m.fc.in_features
    m.fc = nn.Sequential(
        nn.Linear(num_ftrs, 512),
        nn.ReLU(),
        nn.Dropout(0.5),
        nn.Linear(512, num_classes)
    )
    return m


# ────────────────────────────────────────────────────────────────
# LOAD MODELS AT STARTUP
# ────────────────────────────────────────────────────────────────
print("Loading models...")

eff_model = build_resnet50().to(device)
eff_model.load_state_dict(
    torch.load(EFF_WEIGHTS, map_location=device))
eff_model.eval()
print("[SUCCESS] ResNet50 loaded")

dense_model = build_resnet50_subtype(num_classes=4).to(device)
dense_model.load_state_dict(
    torch.load(DENSE_WEIGHTS, map_location=device))
dense_model.eval()
print("[SUCCESS] ResNet50 Subtype model loaded")

# ── Grad-CAM instances ────────────────────────────────────────
eff_gradcam   = GradCAM(eff_model,   model_type="resnet")
dense_gradcam = GradCAM(dense_model, model_type="resnet")

print("[SUCCESS] Grad-CAM ready")

# ── LIME explainers ───────────────────────────────────────────
lime_transform = transforms.Compose([
    transforms.Resize((IMG_SIZE, IMG_SIZE)),
    transforms.ToTensor(),
    transforms.Normalize(mean=[0.485, 0.456, 0.406],
                         std=[0.229, 0.224, 0.225])
])
lime_explainer = LimeExplainer(
    dense_model, lime_transform, device, FOUR_CLASSES)
lime_explainer_binary = LimeExplainer(
    eff_model, lime_transform, device, BINARY_CLASSES)
print("[SUCCESS] LIME explainers ready")

print("Server starting...")


# ────────────────────────────────────────────────────────────────
# IMAGE PREPROCESSING
# ────────────────────────────────────────────────────────────────
def enhance_image(img):
    img    = img.astype(np.uint8)
    img    = cv2.medianBlur(img, 3)
    img    = cv2.equalizeHist(img)
    clahe  = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
    img    = clahe.apply(img)
    kernel = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (3, 3))
    img    = cv2.morphologyEx(img, cv2.MORPH_OPEN,  kernel)
    img    = cv2.morphologyEx(img, cv2.MORPH_CLOSE, kernel)
    return img


def preprocess_image(image_bytes):
    """Convert uploaded image bytes to tensor + return display images"""
    # Load from bytes
    nparr    = np.frombuffer(image_bytes, np.uint8)
    img_orig = cv2.imdecode(nparr, cv2.IMREAD_GRAYSCALE)

    if img_orig is None:
        raise ValueError("Could not read image")

    # Resize and enhance
    img_resized  = cv2.resize(img_orig, (IMG_SIZE, IMG_SIZE))
    img_enhanced = enhance_image(img_resized.copy())

    # Build tensor
    transform = transforms.Compose([
        transforms.ToPILImage(),
        transforms.ToTensor(),
        transforms.Normalize(mean=[0.485, 0.456, 0.406],
                              std=[0.229, 0.224, 0.225])
    ])
    img_rgb    = np.stack([img_enhanced, img_enhanced,
                            img_enhanced], axis=-1)
    img_tensor = transform(img_rgb).unsqueeze(0).to(device)

    return img_tensor, img_resized, img_enhanced


def img_to_base64(img_array, is_color=False):
    """Convert numpy array to base64 string for HTML display"""
    if is_color:
        img_bgr = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
        _, buffer = cv2.imencode('.png', img_bgr)
    else:
        _, buffer = cv2.imencode('.png', img_array)
    return base64.b64encode(buffer).decode('utf-8')


# ────────────────────────────────────────────────────────────────
# ROUTES
# ────────────────────────────────────────────────────────────────
@app.route('/')
def index():
    return render_template('index.html')


@app.route('/analyze/binary', methods=['POST'])
def analyze_binary():
    """Sector 1: Binary classification with Grad-CAM"""
    try:
        if 'image' not in request.files:
            return jsonify({'error': 'No image uploaded'}), 400

        image_file  = request.files['image']
        image_bytes = image_file.read()

        # Preprocess
        img_tensor, img_orig, img_enhanced = \
            preprocess_image(image_bytes)

        # Predict
        with torch.no_grad():
            outputs = eff_model(img_tensor)
            probs   = torch.softmax(outputs, dim=1)[0]

        probs_list  = probs.cpu().numpy()
        pred_idx    = int(np.argmax(probs_list))
        pred_class  = BINARY_CLASSES[pred_idx]
        confidence  = float(probs_list[pred_idx]) * 100

        result = {
            'prediction':   pred_class,
            'confidence':   round(confidence, 1),
            'benign_prob':  round(float(probs_list[0]) * 100, 1),
            'malig_prob':   round(float(probs_list[1]) * 100, 1),
            'is_malignant': pred_idx == 1,
            'original_img': img_to_base64(img_orig),
            'enhanced_img': img_to_base64(img_enhanced),
        }

        # Generate Grad-CAM if malignant
        if pred_idx == 1:
            cam, _ = eff_gradcam.generate(img_tensor, pred_idx)
            overlay, heatmap = generate_heatmap_overlay(
                img_enhanced, cam)
            analysis = analyze_heatmap(cam)

            result['heatmap_img']  = img_to_base64(overlay,
                                                     is_color=True)
            result['cam_analysis'] = analysis

            # Generate binary clinical notes (location-based)
            binary_notes = generate_binary_clinical_notes(
                analysis, confidence)
            result['clinical_notes'] = binary_notes

            size_to_stage = {
                'Small': 'Stage I',
                'Medium': 'Stage II',
                'Large': 'Stage III–IV'
            }
            est_stage = size_to_stage.get(
                analysis['size_label'], 'Unknown')

            result['recommendation'] = (
                f"⚠️ Malignancy detected — "
                f"{analysis['location']} "
                f"({analysis.get('shape', 'Nodular')} mass). "
                f"Estimated {est_stage} based on "
                f"{analysis['activation_pct']}% scan coverage. "
                f"{analysis['spread']}. "
                f"Proceed to Sector 2 for cancer subtype "
                f"identification and full clinical report."
            )

            # ── Dual Grad-CAM (predicted vs counter) ──
            pred_cam, counter_cam, counter_idx = \
                eff_gradcam.generate_dual_cam(
                    img_tensor, pred_idx, num_classes=2)
            counter_overlay, _ = generate_heatmap_overlay(
                img_enhanced, counter_cam)
            result['counter_heatmap'] = img_to_base64(
                counter_overlay, is_color=True)
            result['counter_class'] = BINARY_CLASSES[
                counter_idx]

            # ── AI Explanation Narrative ──
            probs_pct = [float(p) * 100 for p in probs_list]
            result['explanation'] = \
                generate_explanation_narrative(
                    pred_class, confidence, analysis,
                    BINARY_CLASSES, probs_pct,
                    is_binary=True)

        # ── Save to Database ──
        patient_name = request.form.get('patient_name', 'Anonymous').strip()
        if not patient_name:
            patient_name = 'Anonymous'

        diag = Diagnosis(
            patient_name=patient_name,
            sector="Sector 1 (Binary)",
            prediction=pred_class,
            confidence=round(confidence, 1),
            is_malignant=result['is_malignant'],
            downloaded=False
        )

        if result['is_malignant'] and 'cam_analysis' in result:
            diag.stage = est_stage
            diag.location = result['cam_analysis']['location']
            diag.risk_level = 'HIGH RISK'
            diag.activation_pct = result['cam_analysis']['activation_pct']
            diag.num_spots = result['cam_analysis']['num_spots']
            diag.shape = result['cam_analysis']['shape']
            diag.clinical_notes_json = json.dumps(result['clinical_notes'])
        else:
            diag.risk_level = 'LOW RISK'

        db.session.add(diag)
        db.session.commit()
        result['db_id'] = diag.id

        return jsonify(result)

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/analyze/fourclass', methods=['POST'])
def analyze_fourclass():
    """Sector 2: 4-class classification with Grad-CAM + notes"""
    try:
        if 'image' not in request.files:
            return jsonify({'error': 'No image uploaded'}), 400

        image_file  = request.files['image']
        image_bytes = image_file.read()

        # Preprocess
        img_tensor, img_orig, img_enhanced = \
            preprocess_image(image_bytes)

        # Predict
        with torch.no_grad():
            outputs = dense_model(img_tensor)
            probs   = torch.softmax(outputs, dim=1)[0]

        probs_list = probs.cpu().numpy()
        pred_idx   = int(np.argmax(probs_list))
        pred_class = FOUR_CLASSES[pred_idx]
        confidence = float(probs_list[pred_idx]) * 100

        # Grad-CAM
        cam, _ = dense_gradcam.generate(img_tensor, pred_idx)
        overlay, _ = generate_heatmap_overlay(img_enhanced, cam)
        analysis   = analyze_heatmap(cam)

        # Clinical notes
        notes = generate_clinical_notes(
            pred_class, analysis, confidence)

        # ── Dual Grad-CAM (predicted vs counter) ──
        pred_cam, counter_cam, counter_idx = \
            dense_gradcam.generate_dual_cam(
                img_tensor, pred_idx, num_classes=4)
        counter_overlay, _ = generate_heatmap_overlay(
            img_enhanced, counter_cam)

        # ── AI Explanation Narrative ──
        probs_pct = [float(p) * 100 for p in probs_list]
        explanation = generate_explanation_narrative(
            pred_class, confidence, analysis,
            FOUR_CLASSES, probs_pct, is_binary=False)

        result = {
            'prediction':    pred_class,
            'confidence':    round(confidence, 1),
            'probabilities': {
                FOUR_CLASSES[i]: round(float(probs_list[i]) * 100, 1)
                for i in range(4)
            },
            'original_img':     img_to_base64(img_orig),
            'enhanced_img':     img_to_base64(img_enhanced),
            'heatmap_img':      img_to_base64(overlay,
                                                is_color=True),
            'cam_analysis':     analysis,
            'clinical_notes':   notes,
            'counter_heatmap':  img_to_base64(
                counter_overlay, is_color=True),
            'counter_class':    FOUR_CLASSES[counter_idx],
            'explanation':      explanation
        }

        # ── Save to Database ──
        patient_name = request.form.get('patient_name', 'Anonymous').strip()
        if not patient_name:
            patient_name = 'Anonymous'

        diag = Diagnosis(
            patient_name=patient_name,
            sector="Sector 2 (Subtype)",
            prediction=pred_class,
            confidence=round(confidence, 1),
            is_malignant=(pred_class != 'Normal'),
            stage=notes.get('stage', 'N/A') if pred_class != 'Normal' else 'N/A',
            location=analysis['location'],
            risk_level=notes.get('risk', 'ABNORMAL'),
            activation_pct=analysis['activation_pct'],
            num_spots=analysis['num_spots'],
            shape=analysis['shape'],
            clinical_notes_json=json.dumps(notes),
            downloaded=False
        )

        db.session.add(diag)
        db.session.commit()
        result['db_id'] = diag.id

        return jsonify(result)

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/explain/lime', methods=['POST'])
def explain_lime():
    """On-demand LIME explanation for 4-class model"""
    try:
        if 'image' not in request.files:
            return jsonify({'error': 'No image uploaded'}), 400

        image_file  = request.files['image']
        image_bytes = image_file.read()

        # Preprocess
        img_tensor, img_orig, img_enhanced = \
            preprocess_image(image_bytes)

        # Get prediction first
        with torch.no_grad():
            outputs = dense_model(img_tensor)
            probs   = torch.softmax(outputs, dim=1)[0]

        probs_list = probs.cpu().numpy()
        pred_idx   = int(np.argmax(probs_list))
        pred_class = FOUR_CLASSES[pred_idx]

        # Build RGB image for LIME (224x224x3)
        img_rgb = np.stack(
            [img_enhanced, img_enhanced, img_enhanced],
            axis=-1
        ).astype(np.uint8)

        # Run LIME
        pos_overlay, full_overlay, narrative = \
            lime_explainer.explain(
                img_rgb, pred_idx,
                num_samples=100, num_features=8
            )

        if pos_overlay is None:
            return jsonify({'error': narrative}), 500

        result = {
            'lime_positive':  img_to_base64_rgb(pos_overlay),
            'lime_mask':      img_to_base64_rgb(full_overlay),
            'lime_narrative': narrative,
            'prediction':     pred_class
        }

        return jsonify(result)

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/explain/lime-binary', methods=['POST'])
def explain_lime_binary():
    """On-demand LIME explanation for binary model (ResNet50)"""
    try:
        if 'image' not in request.files:
            return jsonify({'error': 'No image uploaded'}), 400

        image_file  = request.files['image']
        image_bytes = image_file.read()

        # Preprocess
        img_tensor, img_orig, img_enhanced = \
            preprocess_image(image_bytes)

        # Get prediction first
        with torch.no_grad():
            outputs = eff_model(img_tensor)
            probs   = torch.softmax(outputs, dim=1)[0]

        probs_list = probs.cpu().numpy()
        pred_idx   = int(np.argmax(probs_list))
        pred_class = BINARY_CLASSES[pred_idx]

        # Build RGB image for LIME (224x224x3)
        img_rgb = np.stack(
            [img_enhanced, img_enhanced, img_enhanced],
            axis=-1
        ).astype(np.uint8)

        # Run LIME
        pos_overlay, full_overlay, narrative = \
            lime_explainer_binary.explain(
                img_rgb, pred_idx,
                num_samples=100, num_features=8
            )

        if pos_overlay is None:
            return jsonify({'error': narrative}), 500

        result = {
            'lime_positive':  img_to_base64_rgb(pos_overlay),
            'lime_mask':      img_to_base64_rgb(full_overlay),
            'lime_narrative': narrative,
            'prediction':     pred_class
        }

        return jsonify(result)

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/chat', methods=['POST'])
def chat_endpoint():
    """Mock AI Chatbot Endpoint"""
    try:
        data = request.get_json()
        if not data or 'message' not in data:
            return jsonify({'error': 'No message provided'}), 400
        
        user_message = data['message'].lower()
        context = data.get('context', {})
        s1 = context.get('s1')
        s2 = context.get('s2')
        
        # Check context first for image analysis questions
        if any(word in user_message for word in ['image', 'scan', 'result', 'tumor', 'stage', 'confidence', 'diagnosis', 'predict']):
            if s2:
                response_text = f"Based on your latest detailed analysis, the AI classified the image as **{s2.get('prediction')}** with **{s2.get('confidence')}%** confidence. "
                notes = s2.get('clinical_notes', {})
                if notes:
                    response_text += f"The tumor is estimated as **{notes.get('stage', 'Unknown Stage')}**. Risk level: **{notes.get('risk', 'Unknown risk')}**."
                return jsonify({'response': response_text})
            elif s1:
                response_text = f"Based on your initial binary analysis, the AI predicted **{s1.get('prediction')}** with **{s1.get('confidence')}%** confidence."
                if s1.get('is_malignant'):
                    cam = s1.get('cam_analysis', {})
                    response_text += f" The mass is located in the **{cam.get('location', 'lung area')}** and is categorized as a **{cam.get('size_label', 'Unknown size')}** mass."
                return jsonify({'response': response_text})

        # Simple medical mock responses
        if 'scan' in user_message or 'ct' in user_message:
            response_text = "A CT (Computed Tomography) scan uses X-rays to create detailed cross-sectional images of the lungs. It is the gold standard for detecting lung nodules and tumors early."
        elif 'treatment' in user_message:
            response_text = "Treatments for lung cancer vary based on the stage and type (e.g., Non-Small Cell vs. Small Cell). Common options include surgery (lobectomy), targeted therapy, chemotherapy, radiation, and immunotherapy."
        elif 'symptom' in user_message:
            response_text = "Common symptoms of lung cancer include a persistent cough, coughing up blood, shortness of breath, chest pain, hoarseness, and unexplained weight loss. If you experience these, please consult a healthcare professional."
        elif 'adenocarcinoma' in user_message:
            response_text = "Adenocarcinoma is the most common type of non-small cell lung cancer (NSCLC). It typically arises in the outer areas of the lungs and is often found before it has spread."
        elif 'squamous' in user_message:
            response_text = "Squamous cell carcinoma is a type of non-small cell lung cancer that usually originates in the central bronchi. It is strongly linked to a history of smoking."
        else:
            response_text = "I am a medical AI assistant. You can ask me about CT scans, lung cancer subtypes, symptoms, and general treatments, or ask me about the results of your recently analyzed image!"

        return jsonify({'response': response_text})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/download/report/<int:db_id>', methods=['GET'])
def download_report(db_id):
    try:
        diag = Diagnosis.query.get(db_id)
        if not diag:
            return jsonify({'error': 'Diagnosis record not found'}), 404

        # Mark as downloaded
        diag.downloaded = True
        db.session.commit()

        # Generate Document
        doc = Document()
        
        # Style margins
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        # Style normal text
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        # Title
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run("LUNG CANCER AI DIAGNOSTIC SYSTEM\nCLINICAL CLINICAL NOTES & PATIENT REPORT")
        run_title.bold = True
        run_title.font.size = Pt(16)
        run_title.font.color.rgb = RGBColor(41, 121, 255) # Custom Accent Blue

        doc.add_paragraph("-" * 80)

        # Patient & Meta Table
        table = doc.add_table(rows=4, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        meta = [
            ("Patient Name:", diag.patient_name),
            ("Report Date:", diag.timestamp.strftime("%Y-%m-%d %H:%M:%S UTC")),
            ("Diagnostic Sector:", diag.sector),
            ("Database Reference ID:", f"LCAI-REF-{diag.id:05d}")
        ]
        for idx, (lbl, val) in enumerate(meta):
            row = table.rows[idx]
            # Label
            cell_lbl = row.cells[0]
            cell_lbl.text = lbl
            cell_lbl.paragraphs[0].runs[0].bold = True
            cell_lbl.paragraphs[0].runs[0].font.size = Pt(10.5)
            # Value
            cell_val = row.cells[1]
            cell_val.text = val
            cell_val.paragraphs[0].runs[0].font.size = Pt(10.5)

        doc.add_paragraph("\n")

        # Diagnosis Heading
        h_diag = doc.add_heading(level=2)
        r_h_diag = h_diag.add_run("1. Diagnostic Summary")
        r_h_diag.bold = True
        r_h_diag.font.size = Pt(14)
        r_h_diag.font.color.rgb = RGBColor(0, 0, 0)

        p_result = doc.add_paragraph()
        p_result.add_run("AI Classification Prediction: ").bold = True
        run_class = p_result.add_run(diag.prediction.upper())
        run_class.bold = True
        run_class.font.size = Pt(12)
        
        # Color coding prediction
        if diag.prediction in ['Malignant', 'Adenocarcinoma', 'Large Cell', 'Squamous Cell']:
            run_class.font.color.rgb = RGBColor(204, 0, 0) # Red
        else:
            run_class.font.color.rgb = RGBColor(0, 153, 76) # Green

        p_conf = doc.add_paragraph()
        p_conf.add_run("Model Confidence Score: ").bold = True
        p_conf.add_run(f"{diag.confidence}%")

        p_risk = doc.add_paragraph()
        p_risk.add_run("Clinical Risk Level: ").bold = True
        r_risk = p_risk.add_run(diag.risk_level or "N/A")
        r_risk.bold = True
        if diag.risk_level in ['HIGH RISK', 'VERY HIGH RISK', 'ABNORMAL']:
            r_risk.font.color.rgb = RGBColor(204, 0, 0)
        else:
            r_risk.font.color.rgb = RGBColor(0, 153, 76)

        if diag.stage and diag.stage != "N/A":
            p_stage = doc.add_paragraph()
            p_stage.add_run("Estimated Nodule Stage: ").bold = True
            p_stage.add_run(diag.stage)

        doc.add_paragraph("\n")

        # Visual Findings Heading
        h_findings = doc.add_heading(level=2)
        r_h_find = h_findings.add_run("2. Explainable AI Nodule Parameters")
        r_h_find.bold = True
        r_h_find.font.size = Pt(14)
        r_h_find.font.color.rgb = RGBColor(0, 0, 0)

        p_find = doc.add_paragraph()
        p_find.add_run("Primary Location: ").bold = True
        p_find.add_run(diag.location or "N/A")

        p_size = doc.add_paragraph()
        p_size.add_run("Nodule Size Class: ").bold = True
        p_size.add_run(f"{diag.activation_pct}% scan area coverage" if diag.activation_pct else "N/A")

        p_shape = doc.add_paragraph()
        p_shape.add_run("Boundary Shape Class: ").bold = True
        p_shape.add_run(diag.shape or "N/A")

        p_spread = doc.add_paragraph()
        p_spread.add_run("Spread Pattern (Foci Count): ").bold = True
        p_spread.add_run("Localized (single focus)" if (diag.num_spots or 0) <= 1 else f"Multiple foci detected ({diag.num_spots})")

        doc.add_paragraph("\n")

        # Clinical Guidance Heading
        h_notes = doc.add_heading(level=2)
        r_h_notes = h_notes.add_run("3. Clinical Guidance & Advisory")
        r_h_notes.bold = True
        r_h_notes.font.size = Pt(14)
        r_h_notes.font.color.rgb = RGBColor(0, 0, 0)

        # Parse notes JSON if exists
        notes_parsed = {}
        if diag.clinical_notes_json:
            try:
                notes_parsed = json.loads(diag.clinical_notes_json)
            except Exception:
                pass

        if notes_parsed:
            if 'location_analysis' in notes_parsed:
                doc.add_paragraph().add_run("Anatomical Pathway Analysis:").bold = True
                p_la = doc.add_paragraph(notes_parsed['location_analysis'])
                p_la.paragraph_format.left_indent = Inches(0.25)
            
            if 'symptoms' in notes_parsed and notes_parsed['symptoms']:
                doc.add_paragraph().add_run("Correlating Clinical Symptoms:").bold = True
                for sym in notes_parsed['symptoms']:
                    p_bullet = doc.add_paragraph(sym, style='List Bullet')
                    p_bullet.paragraph_format.left_indent = Inches(0.25)

            if 'next_steps' in notes_parsed and notes_parsed['next_steps']:
                doc.add_paragraph().add_run("Advisory Diagnostic Protocol (Next Steps):").bold = True
                for step in notes_parsed['next_steps']:
                    p_bullet = doc.add_paragraph(step, style='List Bullet')
                    p_bullet.paragraph_format.left_indent = Inches(0.25)

            if 'treatments' in notes_parsed and notes_parsed['treatments']:
                doc.add_paragraph().add_run("Available Care Protocols & Treatment Options:").bold = True
                for rx in notes_parsed['treatments']:
                    p_bullet = doc.add_paragraph(rx, style='List Bullet')
                    p_bullet.paragraph_format.left_indent = Inches(0.25)
        else:
            p_normal = doc.add_paragraph("No suspicious malignant activation maps found. Diagnostic findings are within normal baseline thresholds.")
            p_normal.paragraph_format.left_indent = Inches(0.25)

        doc.add_paragraph("\n")
        doc.add_paragraph("-" * 80)
        p_disc = doc.add_paragraph("Disclaimer: This report is dynamically generated by an Artificial Intelligence diagnostic decision support system. It represents statistical probabilities of deep learning model features and must be correlated clinically with biopsy confirmation by a board-certified radiologist prior to treatment administration.")
        p_disc.runs[0].font.size = Pt(8.5)
        p_disc.runs[0].italic = True
        p_disc.runs[0].font.color.rgb = RGBColor(128, 128, 128)

        # Write to memory buffer
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        # Format filename
        safe_name = "".join([c if c.isalnum() else "_" for c in diag.patient_name])
        filename = f"LCAI_Report_{safe_name}_{diag.id:04d}.docx"

        return send_file(
            file_stream,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        return jsonify({'error': f'Failed to generate report: {str(e)}'}), 500


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
