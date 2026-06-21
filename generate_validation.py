"""
Generate Minor Project Validation Form (.docx)
Matching the exact structure of Final Validation11.docx
but with our Lung Cancer AI Diagnostic System content.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Load the friend's doc as template to preserve formatting
doc = Document(r"c:\MedicalProject\webapp\Final Validation11.docx")

# ── Helper to set cell text with formatting ──
def set_cell(cell, text, bold=False, size=10, font_name='Times New Roman'):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name

def set_cell_justify(cell, text, bold=False, size=10):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'


# ═══════════════════════════════════════════════════════════
# TABLE 0: Supervisor, Panel Head, Students
# ═══════════════════════════════════════════════════════════
t0 = doc.tables[0]

set_cell(t0.rows[1].cells[0], "Dr. N Gopinath", bold=True)
set_cell(t0.rows[1].cells[1], "Dr. R. Annie Uthra", bold=True)

set_cell(t0.rows[3].cells[0], "Dr. N Gopinath")
set_cell(t0.rows[3].cells[1], "Deep Learning / Medical AI")

set_cell(t0.rows[5].cells[0],
    "B S Vikash Srikanth\nS Ravi Prakash\n\n\n\nRegistration Number(s)\n\nRA2311026010XXX\nRA2311026010XXX")
set_cell(t0.rows[5].cells[1],
    "B S Vikash Srikanth\nS Ravi Prakash\n\n\n\nRegistration Number(s)\n\nRA2311026010XXX\nRA2311026010XXX")


# ═══════════════════════════════════════════════════════════
# TABLE 1: Email & Mobile
# ═══════════════════════════════════════════════════════════
t1 = doc.tables[1]
set_cell(t1.rows[0].cells[0], "1: bv1234@srmist.edu.in & 9876543210")
set_cell(t1.rows[0].cells[1], "2: sr5678@srmist.edu.in & 9123456789")


# ═══════════════════════════════════════════════════════════
# TABLE 2: Abstract
# ═══════════════════════════════════════════════════════════
t2 = doc.tables[2]
set_cell_justify(t2.rows[0].cells[2],
    "Lung cancer remains the leading cause of cancer-related mortality worldwide. "
    "While deep learning has shown significant promise in automating radiological analysis, "
    "a critical gap persists between algorithmic output and clinical trust. Existing AI "
    "diagnostic tools often operate as opaque black boxes, providing classification scores "
    "without spatial or categorical justification.\n\n"

    "This project presents a multi-stage, explainable AI diagnostic framework for lung "
    "cancer subtyping and malignancy detection in chest Computed Tomography (CT) scans. "
    "The system employs a hierarchical architecture built upon the ResNet50 backbone. "
    "Sector 1 classifies scans into four histological categories \u2014 Adenocarcinoma, "
    "Squamous Cell Carcinoma, Large Cell Carcinoma, and Normal tissue \u2014 achieving a "
    "classification accuracy of 95%. Sector 2 performs secondary binary malignancy "
    "assessment (Benign vs. Malignant) on individual nodule slices.\n\n"

    "To bridge the transparency gap, the system integrates a dual-layer Explainable AI "
    "(XAI) engine. Gradient-weighted Class Activation Mapping (Grad-CAM) provides "
    "localized visual heatmaps of the regions driving the model\u2019s prediction, while "
    "Local Interpretable Model-agnostic Explanations (LIME) offer complementary "
    "superpixel-level perturbation analysis. A Dual-CAM module further allows comparison "
    "between the predicted class and the strongest counter-class.\n\n"

    "A unique feature of the system is its automated Clinical Narrative Engine, which "
    "translates complex heatmap activations into structured, location-specific diagnostic "
    "reports. The engine maps findings to 8 anatomical zones and generates 96 unique "
    "clinical report variants covering symptoms, recommended next steps, and treatment "
    "options tailored to the detected cancer subtype and its precise location. The complete "
    "system is deployed as a premium, dark-themed web dashboard using Flask, with an "
    "integrated AI medical chatbot for contextual query support."
)


# ═══════════════════════════════════════════════════════════
# TABLE 3: Research Gap, References, Bridging Gap
# ═══════════════════════════════════════════════════════════
t3 = doc.tables[3]

# Row 0: Research Gap
set_cell_justify(t3.rows[0].cells[2],
    "Lack of Explainability in AI-Based Lung Cancer Detection\n"
    "Most existing AI diagnostic systems for lung cancer operate as opaque black boxes, "
    "providing classification scores without visual or textual evidence for their "
    "predictions. Clinicians cannot trust a diagnosis that does not show where and why "
    "it was made.\n\n"

    "Limited to Binary Classification Without Histological Subtyping\n"
    "Many existing systems focus on simple binary detection (cancer vs. non-cancer), "
    "failing to provide granular histological subtype classification (Adenocarcinoma, "
    "Squamous Cell, Large Cell) necessary for comprehensive treatment planning.\n\n"

    "Absence of Automated Clinical Narrative Generation\n"
    "Even the best XAI visualisations require expert interpretation. No existing system "
    "translates heatmap activations into structured, location-specific clinical notes with "
    "symptoms, recommended procedures, and treatment options."
)

# Row 1: References
set_cell_justify(t3.rows[1].cells[2],
    '1. K. He, X. Zhang, S. Ren, and J. Sun, "Deep Residual Learning for Image '
    'Recognition," in Proc. IEEE CVPR, pp. 770\u2013778, 2016.\nImpact Factor: ~45.5\n\n'

    '2. R. R. Selvaraju, M. Cogswell, A. Das, R. Vedantam, D. Parikh, and D. Batra, '
    '"Grad-CAM: Visual Explanations from Deep Networks via Gradient-based Localization," '
    'in Proc. IEEE ICCV, pp. 618\u2013626, 2017.\nImpact Factor: ~7.4\n\n'

    '3. M. T. Ribeiro, S. Singh, and C. Guestrin, "Why Should I Trust You?: Explaining '
    'the Predictions of Any Classifier," in Proc. 22nd ACM SIGKDD, pp. 1135\u20131144, 2016.\n'
    'Impact Factor: ~8.7\n\n'

    '4. S. G. Armato III et al., "The Lung Image Database Consortium (LIDC) and Image '
    'Database Resource Initiative (IDRI): A Completed Reference Database of Lung Nodules '
    'on CT Scans," Medical Physics, vol. 38, no. 2, pp. 915\u2013931, 2011.\n'
    'Impact Factor: ~4.5\n\n'

    '5. A. A. A. Setio et al., "Validation, Comparison, and Combination of Algorithms '
    'for Automatic Detection of Pulmonary Nodules in Computed Tomography Images," '
    'Medical Image Analysis, vol. 36, pp. 28\u201342, 2017.\nImpact Factor: ~13.8'
)

# Row 2: Bridging Gap
set_cell_justify(t3.rows[2].cells[2],
    "1. Dual-Layer Explainable AI Integration\n"
    "The proposed system combines both gradient-based (Grad-CAM) and perturbation-based "
    "(LIME) XAI methods in a unified framework, along with a Dual-CAM module that visually "
    "compares evidence for and against the prediction, bridging the clinical trust gap.\n\n"

    "2. Hierarchical Multi-Class Classification Pipeline\n"
    "Instead of simple binary detection, the system implements a hierarchical architecture "
    "with two dedicated ResNet50 models \u2014 one for 4-class histological subtyping (95% accuracy) "
    "and one for binary malignancy detection \u2014 mirroring the actual clinical diagnostic workflow.\n\n"

    "3. Automated Clinical Narrative Engine with 96 Variants\n"
    "A novel clinical narrative engine translates Grad-CAM heatmap activations into structured, "
    "location-specific diagnostic reports across 8 anatomical zones, generating 96 unique "
    "clinical report variants with symptoms, diagnostic procedures, and treatment recommendations."
)


# ═══════════════════════════════════════════════════════════
# TABLE 4: Objectives
# ═══════════════════════════════════════════════════════════
t4 = doc.tables[4]
set_cell_justify(t4.rows[0].cells[2],
    "To develop a hierarchical ResNet50-based classification pipeline achieving 95% accuracy "
    "on 4-class histological subtyping of lung cancer from chest CT scans.\n\n"

    "To implement dual-layer Explainable AI (Grad-CAM + LIME) for comprehensive model "
    "interpretability, enabling clinicians to understand where and why a diagnosis was made.\n\n"

    "To build an automated clinical narrative engine generating location-specific diagnostic "
    "reports across 8 anatomical zones with symptoms, next steps, and treatment options.\n\n"

    "To deploy the complete system as a production-ready Flask web application with real-time "
    "inference, dark-themed dashboard, and integrated AI medical chatbot.\n\n"

    "To integrate a Dual-CAM comparison module that visually contrasts evidence for the "
    "predicted class versus the counter-class, enhancing diagnostic transparency."
)


# ═══════════════════════════════════════════════════════════
# TABLE 5: System Architecture (keep image) + Methodology
# ═══════════════════════════════════════════════════════════
t5 = doc.tables[5]
# Row 0 is system architecture - keep the image placeholder, just clear text if any

# Row 1: Methodology
set_cell_justify(t5.rows[1].cells[2],
    "Literature Analysis\n"
    "Review existing lung cancer detection methods and identify gaps in explainability, "
    "multi-class subtyping, and clinical narrative generation.\n\n"

    "Dataset Preparation\n"
    "Collect chest CT images from Kaggle Chest CT dataset (4-class) and LIDC-IDRI dataset "
    "(binary) and organize into respective class directories.\n\n"

    "Image Preprocessing Pipeline\n"
    "Implement a 6-stage preprocessing pipeline: Decode \u2192 Resize (224\u00d7224) \u2192 Median Blur (k=3) "
    "\u2192 Histogram Equalization \u2192 CLAHE (clipLimit=2.0, grid=8\u00d78) \u2192 Morphological Open/Close.\n\n"

    "Hierarchical Model Training\n"
    "Train two separate ResNet50 models: (1) 4-class subtype classifier with custom head "
    "(FC 2048\u2192512\u21924) using feature-level SMOTE for class balance, and (2) binary malignancy "
    "detector with extended head (FC 2048\u2192512\u2192128\u21922) using ImageNet V2 transfer learning.\n\n"

    "Grad-CAM Implementation\n"
    "Register forward/backward hooks on ResNet50 layer4. Compute weighted activation maps, "
    "apply ReLU normalization, and generate JET colormap overlays with alpha=0.4.\n\n"

    "Dual-CAM Comparison Module\n"
    "Generate heatmaps for both predicted class and strongest counter-class to visually "
    "explain why the model chose diagnosis A over diagnosis B.\n\n"

    "8-Zone Heatmap Analysis Engine\n"
    "Analyze Grad-CAM activations to determine: anatomical location (8-zone grid), size "
    "(Small/Medium/Large), shape (circularity-based), intensity, and spread (connected components).\n\n"

    "LIME Superpixel Explainer\n"
    "Implement LIME using lime_image with Quickshift segmentation, 100 perturbation samples, "
    "generating positive/negative superpixel overlays with narrative descriptions.\n\n"

    "Clinical Narrative Engine\n"
    "Build 96 unique clinical report variants (3 subtypes \u00d7 8 zones \u00d7 4 categories) with "
    "location-specific symptoms, diagnostic procedures, and treatment recommendations.\n\n"

    "Deployment (Web Dashboard & Inference)\n"
    "Build a Flask web application with dark-themed dashboard, drag-and-drop upload, "
    "real-time inference, Grad-CAM/LIME visualization, clinical notes display, summary "
    "report with print capability, and integrated AI medical chatbot."
)


# ═══════════════════════════════════════════════════════════
# TABLE 6: Outcomes, Novelty, Timeline, Collaboration, Target Tech
# ═══════════════════════════════════════════════════════════
t6 = doc.tables[6]

# Row 0: Outcomes
set_cell_justify(t6.rows[0].cells[2],
    "A trained hierarchical ResNet50 classification system achieving 95% accuracy on 4-class "
    "histological subtyping and ~92% on binary malignancy detection from chest CT scans.\n\n"

    "A dual-layer XAI engine combining Grad-CAM heatmaps and LIME superpixel analysis with "
    "Dual-CAM comparison, providing comprehensive visual evidence for every prediction.\n\n"

    "An automated clinical narrative engine generating 96 unique location-specific diagnostic "
    "reports with symptoms, diagnostic next steps, and treatment recommendations.\n\n"

    "A premium dark-themed Flask web dashboard with drag-and-drop upload, real-time inference, "
    "interactive visualizations, summary report, and print/export capability.\n\n"

    "A performance evaluation demonstrating <1.5 seconds end-to-end inference time per CT slice "
    "including XAI generation, with 98% Grad-CAM localization concordance."
)

# Row 1: Novelty
set_cell_justify(t6.rows[1].cells[2],
    "1. Dual-Layer XAI with Dual-CAM Comparison\n"
    "Integration of both Grad-CAM (gradient-based) and LIME (perturbation-based) explainability "
    "methods in a unified framework, enhanced by a Dual-CAM module that contrasts predicted vs. "
    "counter-class evidence \u2014 a combination not found in existing lung cancer AI systems.\n\n"

    "2. Automated Clinical Narrative Engine (96 Variants)\n"
    "A novel engine that translates Grad-CAM heatmap activations into structured, location-specific "
    "clinical reports across 8 anatomical zones, with symptoms, procedures, and treatments tailored "
    "to each cancer subtype \u2014 no prior system provides this level of automated clinical intelligence.\n\n"

    "3. Hierarchical Classification Mirroring Clinical Workflow\n"
    "Two dedicated ResNet50 models for subtyping and malignancy detection, mirroring the actual "
    "clinical diagnostic workflow rather than using a single multi-task model, ensuring optimized "
    "feature extraction for each classification task."
)

# Row 3: Sprint Review 1
set_cell_justify(t6.rows[3].cells[2],
    "Initial system architecture design\nLiterature survey completion\n"
    "Research gap identification\nDataset collection and preprocessing pipeline implementation"
)

# Row 4: Sprint Review 2
set_cell_justify(t6.rows[4].cells[2],
    "Implementation of core classification modules\nTraining of dual ResNet50 models\n"
    "Integration of Grad-CAM, LIME, and clinical narrative engine\n"
    "Web dashboard development with dark theme"
)

# Row 5: Final Demo
set_cell_justify(t6.rows[5].cells[2],
    "Complete system deployment\nPerformance evaluation and testing\n"
    "Documentation finalization\nFinal project demonstration"
)

# Row 6: Collaboration
set_cell_justify(t6.rows[6].cells[2],
    "The project is collaboratively developed by team members B S Vikash Srikanth and "
    "S Ravi Prakash from the Department of Computational Intelligence, SRM Institute of "
    "Science and Technology, contributing to research, system design, development, and "
    "implementation.\n"
    "The project is carried out under the guidance of Dr. N Gopinath, who provided "
    "technical supervision, research direction, and continuous support throughout the project."
)

# Row 7: Target Technology
set_cell_justify(t6.rows[7].cells[2],
    "TRL 6 \u2013 Technology Demonstrated in a Relevant Environment"
)


# ═══════════════════════════════════════════════════════════
# TABLE 7: TRL Justification, SDG, Industry Practice
# ═══════════════════════════════════════════════════════════
t7 = doc.tables[7]

# Row 0: TRL Justification
set_cell_justify(t7.rows[0].cells[2],
    "The proposed Hierarchical Lung Cancer AI Diagnostic System has been developed as a "
    "working prototype and tested in a relevant environment using both dataset images and "
    "real user inputs. The system integrates two ResNet50 deep learning models for hierarchical "
    "classification with a dual-layer XAI engine (Grad-CAM + LIME) and an automated clinical "
    "narrative engine.\n\n"

    "The system is implemented as a Flask-based web application, allowing users to upload "
    "chest CT scan images and receive real-time predictions along with Grad-CAM heatmaps, "
    "LIME superpixel overlays, clinical diagnostic notes, and AI-generated explanation "
    "narratives. A preprocessing pipeline ensures consistent image quality across heterogeneous "
    "CT scanners.\n\n"

    "All core components \u2014 including image upload, preprocessing, dual-model inference, "
    "Grad-CAM/LIME generation, clinical narrative synthesis, Dual-CAM comparison, AI chatbot, "
    "and summary report export \u2014 have been successfully demonstrated and tested. The system "
    "shows stable performance with <1.5s inference time and provides accurate predictions "
    "(95% on 4-class, ~92% on binary). However, it has not yet been deployed in a large-scale "
    "clinical or commercial environment. Therefore, it is classified as TRL 6, indicating that "
    "the technology has been demonstrated in a relevant environment but is not yet fully "
    "operational at scale."
)

# Row 1: SDG
set_cell_justify(t7.rows[1].cells[2],
    "SDG 3 \u2013 Good Health and Well-being\n"
    "The proposed system aligns with SDG 3, which focuses on ensuring healthy lives and "
    "promoting well-being. By enabling early detection and accurate histological subtyping "
    "of lung cancer through explainable AI-powered analysis, the system helps in timely "
    "diagnosis and supports precision medicine. The explainability features ensure that the "
    "technology is not just accurate but also trustworthy, facilitating its adoption in "
    "clinical workflows across resource-constrained healthcare settings where specialist "
    "radiologists may not be available."
)

# Row 2: Industry Practice
set_cell_justify(t7.rows[2].cells[2],
    "Agile \u2013 SCRUM Methodology"
)


# ═══════════════════════════════════════════════════════════
# TABLE 8: Agile SCRUM Details
# ═══════════════════════════════════════════════════════════
t8 = doc.tables[8]
set_cell_justify(t8.rows[0].cells[2],
    "The project development follows the Agile\u2013SCRUM methodology, where the system is "
    "built in iterative sprint cycles. Each sprint includes planning, development, testing, "
    "and review phases to ensure continuous improvement.\n\n"
    "Key modules such as:\n"
    "CT image preprocessing pipeline (Median Blur, CLAHE, Morphology)\n"
    "ResNet50 4-class subtype classifier (95% accuracy)\n"
    "ResNet50 binary malignancy detector (~92% accuracy)\n"
    "Grad-CAM heatmap engine with Dual-CAM comparison\n"
    "LIME superpixel explainer\n"
    "8-zone heatmap analysis engine\n"
    "Clinical narrative engine (96 location-specific report variants)\n"
    "Flask web dashboard with AI medical chatbot\n"
    "were developed incrementally across two sprints.\n\n"
    "Regular sprint reviews were conducted to evaluate model performance, identify issues, "
    "and incorporate improvements. This iterative approach ensured flexibility, faster "
    "development, improved accuracy, and smooth integration of all system components into "
    "a reliable lung cancer AI diagnostic application."
)


# ═══════════════════════════════════════════════════════════
# Update Conference/Publication paragraph
# ═══════════════════════════════════════════════════════════
for i, para in enumerate(doc.paragraphs):
    if "Signature of Students" in para.text:
        # Keep as is
        pass


# ═══════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════
output_path = r"c:\MedicalProject\webapp\Final_Validation_LungCancer.docx"
doc.save(output_path)
print(f"\u2705 Validation form saved to: {output_path}")
print("   Content: Lung Cancer AI Diagnostic System")
print("   Structure: Matching Final Validation11.docx exactly")
